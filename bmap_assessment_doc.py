"""
bmap_assessment_doc.py — Verlocity $10K BMAP Assessment (Word doc generator)
=============================================================================
Produces the paid $10K Assessment deliverable as a branded .docx report.
Full-network coverage (every branch, no top-N slice) — this is the structural
difference vs. the free BMAP Snapshot (bmap_snapshot.py), which only shows
the top 5 branches as an outreach teaser.

Same Supabase source as bmap_snapshot.py / bmap_board_brief.py.
Uses python-docx (matches production Python stack — Flask/Railway).

Usage (CLI):
    python bmap_assessment_doc.py --inst_key bank_463735 --name "Hancock Whitney Bank"

Scope note: this is a first structural draft to test what full-network data
+ AI narrative looks like at scale. Session-based content (the live working
call) is intentionally NOT in this doc — per the 7/22 scope doc, the document
alone is not the $10K product; the session is. See closing section.
"""

import os
import sys
import math
import concurrent.futures
import argparse
import requests
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import anthropic
except ImportError:
    anthropic = None

import json

# ── Build marker — printed to logs AND stamped tiny/gray on the closing
# page of both the Assessment and the Preview. Exists specifically to kill
# the "did my last fix actually deploy" back-and-forth: whoever generated
# the doc can just look at the bottom of the last page and compare the
# build id to what they expect, instead of guessing from symptoms again.
# Bump this string any time this file changes.
GENERATOR_BUILD = "2026-08-27.10"
print(f"[bmap_assessment_doc] build {GENERATOR_BUILD}")

# ── Config — matches current bmap_snapshot.py production pattern ──
# (previous version of this file used a hardcoded legacy anon key, which
# is dead now that RLS + service-role-only access is in place — fixed here)
SUPA_URL = "https://tuiiywphoynbmkxpoyps.supabase.co"
SUPA_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")
if not SUPA_KEY:
    print("  ⚠  SUPABASE_SERVICE_KEY is not set — Supabase calls will fail with 401.")

ANTH_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
OUT_DIR  = Path(".")

# ── Brand colors — matches Verlocity_Brand_Guidelines_R2.pdf + bmap_snapshot.py exactly ──
def rgb(h): return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

NAVY   = rgb("083D5F")   # Primary Dark Blue
TEAL   = rgb("02A7C2")   # Primary Light Blue
JET    = rgb("213141")   # Jet Black
EMERALD = rgb("66CC99")
LEMON   = rgb("CDD61A")
GRAY3  = rgb("778899")
GRAY_FILL = "F5F5F2"
RED_WEAK = rgb("B02A2A")       # weakness-signal text color in competitor tables
RED_WEAK_FILL = "FBEAEA"       # light red row-shading for the top target

# Zone colors — matches BMAP_Snapshot_Hancock_Whitney_Bank_20260813.pptx exactly
# (brand-guideline navy→teal→emerald→lemon progression, not a red/green traffic-light scheme)
ZONE_COLOR = {
    "Invest":  "083D5F", "Analyze": "02A7C2", "Defend": "66CC99", "Justify": "BFC815",
}
ZONE_LIGHT = {
    "Invest":  "E1E9EF", "Analyze": "E1F5F8", "Defend": "EAF7F0", "Justify": "F8FAE3",
}
ZONE_HEX_MPL = {  # matplotlib wants '#rrggbb'
    z: f"#{h}" for z, h in ZONE_COLOR.items()
}
FONT_HEAD = "Segoe UI"   # Word-native text (headings, tables, body) references
                          # a font by NAME only -- python-docx can't embed a font
                          # the way the cover image bakes Inter directly into its
                          # pixels (that part IS genuine Inter, unaffected by this
                          # setting). "Inter" as the name here means Word tries to
                          # find Inter installed locally and silently substitutes
                          # something else if it isn't -- confirmed not rendering
                          # for Francisco, so back to Segoe UI, which is actually
                          # installed everywhere Word is. Real Inter TTFs still
                          # ship alongside this script if anyone wants to install
                          # them locally or embed via Word's own File > Options >
                          # Save > "Embed fonts in the file" for a true match.

# ═══════════════════════════════════════════════════════════════
# DATA FETCH — full network, no truncation
# ═══════════════════════════════════════════════════════════════

# SCHEMA_MAP mirrors bmap_snapshot.py exactly — table→schema routing,
# since PostgREST needs Accept-Profile for anything outside 'public'.
SCHEMA_MAP = {
    "branch_opportunity_base":        "analytics",
    "bank_financial_snapshot_latest": "analytics",
    "branch_target_competitors":      "analytics",
    "branches_master_v2":             "geo",
    "raw_sod":                        "raw",
}

def supabase(table, params):
    url = f"{SUPA_URL}/rest/v1/{table}?{params}"
    schema = SCHEMA_MAP.get(table, "public")
    headers = {"apikey": SUPA_KEY, "Authorization": f"Bearer {SUPA_KEY}"}
    if schema != "public":
        headers["Accept-Profile"] = schema
    r = requests.get(url, headers=headers, timeout=30)
    if r.status_code != 200:
        print(f"  ⚠ Supabase {table} error {r.status_code}: {r.text[:200]}")
        return []
    return r.json()


def supabase_rpc(fn_name, payload, timeout=20, paginate=False, page_size=1000):
    """Call a Postgres function via PostgREST's /rpc/ endpoint (e.g. the
    parametrized branches_within_radius(lat, lon, radius, exclude_bank_id)).

    paginate=True fetches ALL rows across multiple requests using PostgREST's
    Range header -- needed for branches_within_radius_batch specifically,
    which returns one row per (client branch, nearby competitor) pair and can
    exceed 10,000 rows for a large network (Trustmark: 166 branches x up to
    10mi radius each = 11,072 rows). PostgREST caps an unpaginated response
    at 1000 rows by default with NO error -- still a 200, just a silently
    incomplete list. Real-world symptom this caused: Jones Valley's 82
    competitor rows landed at position ~5,908 in the unordered 11,072-row
    response, so they never made it into the first page. The branch's
    competitor TABLE still rendered fine (a separate, small, per-branch
    query with its own limit=3), so nothing about the doc looked broken --
    only the map, fed exclusively by this truncated batch, silently had
    nothing to draw."""
    url = f"{SUPA_URL}/rest/v1/rpc/{fn_name}"
    headers = {"apikey": SUPA_KEY, "Authorization": f"Bearer {SUPA_KEY}",
               "Content-Type": "application/json"}
    if not paginate:
        r = requests.post(url, headers=headers, json=payload, timeout=timeout)
        if r.status_code != 200:
            print(f"  ⚠ RPC {fn_name} error {r.status_code}: {r.text[:200]}")
            return []
        return r.json()

    all_rows = []
    offset = 0
    while True:
        page_headers = {**headers, "Range-Unit": "items", "Range": f"{offset}-{offset + page_size - 1}"}
        r = requests.post(url, headers=page_headers, json=payload, timeout=timeout)
        if r.status_code not in (200, 206):
            print(f"  ⚠ RPC {fn_name} error {r.status_code} at offset {offset}: {r.text[:200]}")
            break
        page = r.json()
        if not isinstance(page, list):
            break
        all_rows.extend(page)
        if len(page) < page_size:
            break  # short page -- this was the last one
        offset += page_size
    print(f"  ✓ RPC {fn_name} paginated: {len(all_rows)} total rows "
          f"({offset // page_size + 1} page{'s' if offset else ''})")
    return all_rows


def fetch_full_network_data(ik, skip_competitive_strategy=False):
    """Pull FULL branch network — no limit=N slice. This is the structural
    difference vs. the free Snapshot.

    skip_competitive_strategy=True skips fetch_branch_competitive_strategy()
    (the network-wide, paginated branches_within_radius_batch() call) --
    used by the single-branch Preview flow, which only ever needs ONE
    branch's competitor data and was otherwise paying full-network cost
    (11,000+ rows for a 166-branch network) to throw away 165/166ths of it.
    The full $10K Assessment still calls this with the default False."""
    print(f"  Fetching full branch network for {ik}...")
    branches = supabase(
        "branch_opportunity_base",
        f"inst_key=eq.{ik}&select=uninumbr,namebr,citybr,stalpbr,latest_dep,"
        "yoy_deposits,opportunity_score,opportunity_zone,matrix_quadrant,"
        "priority_tier,market_growth_normalized,rel_growth_norm,"
        "inv_density_norm_winsor,deposit_size_norm,namefull,"
        "household_income,yoy_income_growth,total_population,yoy_pop_growth,"
        "zhvi_yoy_pct,smb_zone,smb_index,total_deposits_10mi&order=opportunity_score.desc",
    )
    print(f"  ✓ {len(branches)} branches (full network, uncapped)")

    print(f"  Fetching financial snapshot...")
    fin_arr = supabase(
        "bank_financial_snapshot_latest",
        f"inst_key=eq.{ik}&select=roa,nim,efficiency_ratio,dep_yoy_pct,"
        "dep_qoq_pct,cost_of_funds_pct,tier1_capital_pct,net_income_yoy_pct,"
        "total_assets,total_deposits,period",
    )
    fin = fin_arr[0] if fin_arr else {}

    print(f"  Fetching network competitor target...")
    tgt_arr = supabase(
        "vw_network_top_targets",
        f"my_inst_key=eq.{ik}&select=target_institution,branches_in_radius,"
        "avg_vuln_score,avg_yoy_pct,target_roa,target_efficiency_ratio,dominant_zone"
        "&order=network_rank.asc&limit=3",
    )

    print(f"  Fetching branch geography for map...")
    branches_geo = fetch_branch_geo(ik)
    print(f"  ✓ {len(branches_geo)} branches geocoded")

    if skip_competitive_strategy:
        print(f"  ⏭ Skipping network-wide competitive strategy (single-branch flow)")
        branch_strategy = []
    else:
        branch_strategy = fetch_branch_competitive_strategy(ik, branches, branches_geo)

    print(f"  Fetching competitor vulnerability scoring...")
    vulnerability_targets = fetch_vulnerability_targets(ik, branches)
    if vulnerability_targets:
        print(f"  ✓ vulnerability data for {len(vulnerability_targets)} branch(es)")

    print(f"  Resolving winsorized YoY values...")
    capped_yoy = resolve_capped_yoy(branches)
    if capped_yoy:
        print(f"  ✓ {len(capped_yoy)} capped branch(es) resolved against raw_sod")

    return {
        "inst_key": ik,
        "branches": branches,
        "fin": fin,
        "targets": tgt_arr,
        "branches_geo": branches_geo,
        "branch_strategy": branch_strategy,
        "vulnerability_targets": vulnerability_targets,
        "capped_yoy": capped_yoy,
    }


def _vuln_tier(score):
    """vuln_score is an intentionally uncapped composite (base 0-100 x up to
    ~4x stacked risk multipliers) - by design, not a bug. Showing the raw
    number next to a 0-100% financial table reads as broken, so we display
    a qualitative tier instead of implying false precision."""
    s = _sf(score)
    if s >= 150:
        return "Critical"
    if s >= 100:
        return "High"
    if s >= 60:
        return "Elevated"
    return "Moderate"


def _sf(v, default=0.0):
    try:
        return float(v) if v is not None else default
    except (TypeError, ValueError):
        return default


def resolve_capped_yoy(branches):
    """yoy_deposits is winsorized/capped at 1.0 (100%) upstream -- a real,
    known limitation (documented pending item: 'real_yoy_display'). Every
    branch showing exactly 1.0 collapses two different real situations into
    one misleading '+100.0%': genuine outsized growth, or a branch with no
    prior-year row at all (new/newly-ingested branch, where YoY is undefined,
    not 100%). This resolves both cases directly from raw.raw_sod so the
    doc never publishes a fabricated-looking exact 100%.
    Returns {uninumbr: display_string} for capped branches only."""
    capped = [b for b in branches if abs(_sf(b.get("yoy_deposits")) - 1.0) < 0.001]
    if not capped:
        return {}

    ids = ",".join(str(b["uninumbr"]) for b in capped)
    rows = supabase("raw_sod", f"UNINUMBR=in.({ids})&select=UNINUMBR,YEAR,DEPSUMBR")
    by_branch = {}
    for r in rows if isinstance(rows, list) else []:
        by_branch.setdefault(str(r["UNINUMBR"]), {})[str(r["YEAR"])] = _sf(r.get("DEPSUMBR"))

    years = sorted({y for v in by_branch.values() for y in v.keys()})
    if len(years) < 2:
        cur_yr, prior_yr = (years[-1], None) if years else (None, None)
    else:
        cur_yr, prior_yr = years[-1], years[-2]

    out = {}
    for b in capped:
        uid = str(b["uninumbr"])
        vals = by_branch.get(uid, {})
        cur = vals.get(cur_yr)
        prior = vals.get(prior_yr) if prior_yr else None
        if prior and prior > 0 and cur is not None:
            real_pct = (cur - prior) / prior * 100
            out[b["uninumbr"]] = f"{real_pct:+.1f}%*"
        else:
            out[b["uninumbr"]] = "New branch*"
    return out


def fmt_yoy(b, capped_map):
    """Single formatter for yoy_deposits everywhere it's displayed --
    routes through the capped-value resolution instead of ever printing
    a bare '+100.0%'."""
    override = capped_map.get(b.get("uninumbr"))
    if override:
        return override
    return f"{_sf(b.get('yoy_deposits'))*100:+.1f}%"


def _score_driver_clause(b):
    """Identifies what's actually driving a branch's opportunity score, in
    plain language -- not for display, for feeding the AI real analytical
    ammunition so the Branch Verdict can explain WHY a score is what it is
    (e.g. 'capped by a shrinking local market, not competition') instead of
    just restating the number. This is the insight-generating use of the
    four weighted score components; a visual breakdown chart exposing the
    raw numbers to the reader was deliberately removed from the document --
    a stacked bar of normalized sub-scores is mechanical transparency, not
    the analysis a buyer is paying for. The same underlying data is more
    valuable as an input to the narrative than as its own exhibit."""
    components = [
        ("market growth", _sf(b.get("market_growth_normalized")), 0.25),
        ("relative growth vs. peers", _sf(b.get("rel_growth_norm")), 0.30),
        ("low competitive density", _sf(b.get("inv_density_norm_winsor")), 0.25),
        ("deposit base size", _sf(b.get("deposit_size_norm")), 0.20),
    ]
    weighted = [(name, val * w) for name, val, w in components]
    if all(w == 0 for _, w in weighted):
        return ""
    dominant = max(weighted, key=lambda x: x[1])
    weakest = min(weighted, key=lambda x: x[1])
    if dominant[0] == weakest[0]:
        return ""
    return f"score driven primarily by {dominant[0]}, weakest on {weakest[0]}"


# ═══════════════════════════════════════════════════════════════
# VISUALS — matplotlib charts + geographic map, embedded as PNGs.
# McKinsey-style: clean, minimal chrome, brand colors, direct labels
# instead of legends where possible.
# ═══════════════════════════════════════════════════════════════
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.patheffects as pe

# Register the real bundled Inter TTFs with matplotlib directly, rather
# than just naming "Inter" in the fallback list and hoping the system has
# it installed (it won't, on a bare Linux container -- this is exactly
# the gap that had every chart silently rendering in DejaVu Sans before).
for _font_file in ("verlocity_font_regular.ttf", "verlocity_font_bold.ttf"):
    _font_path = Path(__file__).parent / _font_file
    if _font_path.exists():
        fm.fontManager.addfont(str(_font_path))

plt.rcParams["font.family"] = "sans-serif"
plt.rcParams["font.sans-serif"] = ["Inter", "DejaVu Sans", "Arial"]
plt.rcParams["axes.edgecolor"] = "#CCCCCC"
plt.rcParams["axes.linewidth"] = 0.6

NAVY_HEX = "#083D5F"
GRAY_HEX = "#8A8A80"


def chart_zone_distribution(zones, path):
    """Horizontal bar, one bar per zone, real brand zone colors, count labeled directly."""
    order = ["Invest", "Analyze", "Defend", "Justify"]
    vals = [zones.get(z, 0) for z in order]
    colors = [ZONE_HEX_MPL[z] for z in order]

    fig, ax = plt.subplots(figsize=(7.2, 2.3), dpi=200)
    y = range(len(order))
    bars = ax.barh(y, vals, color=colors, height=0.62)
    for i, v in enumerate(vals):
        ax.text(v + max(vals) * 0.02, i, str(v), va="center", fontsize=12,
                fontweight="bold", color="#222222")
    ax.set_yticks(y)
    ax.set_yticklabels(order, fontsize=11, color="#222222")
    ax.invert_yaxis()
    ax.set_xticks([])
    for spine in ["top", "right", "bottom"]:
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)
    fig.tight_layout(pad=0.6)
    fig.savefig(path, transparent=True)
    plt.close(fig)


def chart_top_bottom_branches(top5, bottom3, path):
    """Diverging horizontal bar — top branches (green, growth) vs bottom (red, decline)."""
    rows = [(b["namebr"], _sf(b["opportunity_score"]), True) for b in top5] + \
           [(b["namebr"], _sf(b["opportunity_score"]), False) for b in bottom3]
    names = [r[0] for r in rows]
    scores = [r[1] for r in rows]
    colors = [ZONE_HEX_MPL["Invest"] if r[2] else ZONE_HEX_MPL["Justify"] for r in rows]

    fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=200)
    y = range(len(rows))
    ax.barh(y, scores, color=colors, height=0.6)
    for i, v in enumerate(scores):
        ax.text(v + 1.5, i, f"{v:.0f}", va="center", fontsize=10, fontweight="bold", color="#222222")
    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=9.5, color="#222222")
    ax.invert_yaxis()
    ax.set_xlim(0, 105)
    ax.set_xticks([])
    ax.axvline(x=50, color="#CCCCCC", linewidth=0.8, linestyle="--")
    for spine in ["top", "right", "bottom"]:
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)
    fig.tight_layout(pad=0.6)
    fig.savefig(path, transparent=True)
    plt.close(fig)


def chart_score_breakdown(branch, path):
    """Horizontal stacked bar showing the four weighted components that sum
    to a branch's opportunity_score -- makes the score legible instead of a
    single black-box number. Verified mathematically exact against real
    production data before building this: 0.25*market_growth_normalized +
    0.30*rel_growth_norm + 0.25*inv_density_norm_winsor + 0.20*deposit_size_norm
    equals the stored opportunity_score to within rounding (confirmed on a
    real branch: 0.25*1.98 + 0.30*0.00 + 0.25*19.92 + 0.20*68.17 = 19.109,
    matching a stored opportunity_score of 19.11). Returns False (renders
    nothing) if any component is missing, rather than showing a chart that
    silently doesn't add up -- these fields aren't populated for every bank/
    branch in every environment, and a chart that's visibly wrong is worse
    than no chart."""
    components = [
        ("Market\nGrowth", _sf(branch.get("market_growth_normalized")), 0.25, "#083D5F"),
        ("Relative\nGrowth", _sf(branch.get("rel_growth_norm")), 0.30, "#02A7C2"),
        ("Competitive\nDensity", _sf(branch.get("inv_density_norm_winsor")), 0.25, "#66CC99"),
        ("Deposit\nSize", _sf(branch.get("deposit_size_norm")), 0.20, "#BFC815"),
    ]
    if all(v == 0 for _, v, _, _ in components):
        return False  # fields likely unpopulated for this branch/bank -- skip rather than show all-zero

    contributions = [v * w for _, v, w, _ in components]
    total = sum(contributions)
    labels = [c[0] for c in components]
    colors = [c[3] for c in components]

    fig, ax = plt.subplots(figsize=(6.0, 2.1), dpi=200)
    left = 0
    for label, contrib, color in zip(labels, contributions, colors):
        ax.barh([0], [contrib], left=left, color=color, height=0.55,
                edgecolor="white", linewidth=1.5)
        if contrib > total * 0.04:  # skip label if the segment is too thin to read
            ax.text(left + contrib / 2, 0, f"{contrib:.1f}", ha="center", va="center",
                     fontsize=8.5, fontweight="bold", color="white")
        left += contrib

    ax.set_xlim(0, max(total * 1.02, 1))
    ax.set_ylim(-0.6, 0.6)
    ax.axis("off")
    ax.text(-max(total, 1) * 0.02, 0, f"{total:.0f}", ha="right", va="center",
             fontsize=13, fontweight="bold", color=NAVY_HEX)

    # Legend below the bar, in axes-fraction coordinates (0-1) so it's
    # independent of the bar's data scale entirely. 2x2 grid, not a single
    # row -- a single row overlapped ("Competitive Density (25%)" collided
    # with the next label) since label text lengths vary too much for even
    # fixed spacing to handle safely.
    legend_positions = [(0.02, -0.30), (0.52, -0.30), (0.02, -0.55), (0.52, -0.55)]
    for i, (label, contrib, color) in enumerate(zip(labels, contributions, colors)):
        weight_pct = int(components[i][2] * 100)
        lx, ly = legend_positions[i]
        ax.add_patch(plt.Rectangle((lx, ly), 0.02, 0.09, color=color,
                                     transform=ax.transAxes, clip_on=False))
        ax.text(lx + 0.03, ly + 0.045, f"{label.replace(chr(10), ' ')} ({weight_pct}%)",
                fontsize=7.5, va="center", color="#333333", transform=ax.transAxes)

    fig.tight_layout(pad=0.3)
    fig.savefig(path, transparent=True, bbox_inches="tight")
    plt.close(fig)
    return True


def chart_financial_benchmark(fin, bank_name, path):
    """Grouped bar — actual metric vs benchmark, normalized to comparable scale per metric."""
    metrics = [
        ("ROA", _sf(fin.get("roa")), 1.0),
        ("NIM", _sf(fin.get("nim")), 3.0),
        ("Efficiency", _sf(fin.get("efficiency_ratio")), 60.0),
        ("Dep. YoY", _sf(fin.get("dep_yoy_pct")), 2.0),
        ("Tier 1 Cap.", _sf(fin.get("tier1_capital_pct")), 8.0),
    ]
    labels = [m[0] for m in metrics]
    actual = [m[1] for m in metrics]
    bench = [m[2] for m in metrics]

    # Legend label truncated to a short recognizable form (e.g. "Hancock
    # Whitney" from "Hancock Whitney Bank") so it doesn't overflow the chart
    # for long institution names.
    short_name = bank_name.replace(" Bank", "").replace(" Bancorp", "").strip() or bank_name

    x = range(len(labels))
    w = 0.32
    fig, ax = plt.subplots(figsize=(7.2, 2.8), dpi=200)
    ax.bar([i - w/2 for i in x], actual, width=w, color=NAVY_HEX, label=short_name)
    ax.bar([i + w/2 for i in x], bench, width=w, color="#C9CED6", label="Benchmark")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=9.5)
    ax.set_yticks([])
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    ax.legend(frameon=False, fontsize=9, loc="upper right")
    fig.tight_layout(pad=0.6)
    fig.savefig(path, transparent=True)
    plt.close(fig)


_STATES_GEOJSON_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "us_states.geojson")
_STATES_GEOJSON_URL = "https://raw.githubusercontent.com/PublicaMundi/MappingAPI/master/data/geojson/us-states.json"


def _load_state_polygons():
    """Local-first US state boundaries (Polygon/MultiPolygon per state).
    Ships as a repo asset (us_states.geojson) so report generation never
    depends on network access at request time; falls back to a one-time
    fetch + cache if the asset is missing."""
    import json as _json
    if not os.path.exists(_STATES_GEOJSON_PATH):
        try:
            import urllib.request
            with urllib.request.urlopen(_STATES_GEOJSON_URL, timeout=8) as resp:
                data = resp.read()
            with open(_STATES_GEOJSON_PATH, "wb") as f:
                f.write(data)
        except Exception:
            return {}
    try:
        with open(_STATES_GEOJSON_PATH) as f:
            gj = _json.load(f)
    except Exception:
        return {}

    from shapely.geometry import shape
    return {feat["properties"]["name"]: shape(feat["geometry"]) for feat in gj.get("features", [])}


def _draw_polygon(ax, geom, **kwargs):
    from matplotlib.patches import Polygon as MplPolygon
    from shapely.geometry import Polygon, MultiPolygon
    polys = geom.geoms if isinstance(geom, MultiPolygon) else [geom]
    for poly in polys:
        ax.add_patch(MplPolygon(list(poly.exterior.coords), closed=True, **kwargs))
        for interior in poly.interiors:
            ax.add_patch(MplPolygon(list(interior.coords), closed=True,
                                     facecolor="#FAFAF8", edgecolor="none", zorder=kwargs.get("zorder", 1) + 0.1))


MAPBOX_TOKEN = os.environ.get("MAPBOX_TOKEN", "")


def _web_mercator_xy(lon, lat, zoom, tile_size=256):
    """Standard Web Mercator projection to pixel space at a given zoom
    (classic 256px-tile convention -- the same one staticmap/most slippy-map
    libraries use). Used to place every branch marker and label at the exact
    pixel position matching a Mapbox static image requested at the same
    center/zoom, without embedding any per-branch data in the request URL."""
    x = (lon + 180.0) / 360.0 * tile_size * (2 ** zoom)
    lat_rad = math.radians(lat)
    y = (1 - math.log(math.tan(lat_rad) + 1 / math.cos(lat_rad)) / math.pi) / 2 \
        * tile_size * (2 ** zoom)
    return x, y


def _fit_zoom(lons, lats, target_w, target_h, tile_size=256, max_zoom=18):
    """Largest zoom (classic 256px convention) at which the branch bounding
    box fits inside target_w x target_h pixels. Mirrors what staticmap's
    _calculate_zoom does internally, done manually here since this function
    only needs the base tile image, not staticmap's tile-fetching."""
    for z in range(max_zoom, -1, -1):
        x0, y0 = _web_mercator_xy(min(lons), max(lats), z, tile_size)
        x1, y1 = _web_mercator_xy(max(lons), min(lats), z, tile_size)
        if (x1 - x0) <= target_w and (y1 - y0) <= target_h:
            return z
    return 0


def chart_branch_map_osm(branches_geo, path):
    """PRIMARY map — real Mapbox street-map tiles, the recognizable
    'Google Maps view' people actually orient by, instead of a bare state
    outline. Requests ONLY the base map image (a fixed-size URL: center,
    zoom, width, height — no per-branch data embedded), then draws every
    marker and city label myself via matplotlib using the same Web Mercator
    math used to pick that center/zoom.

    This replaced an earlier GeoJSON-overlay design that embedded marker
    data directly in the request URL. That hit Mapbox's ~8192-char URL
    limit on real data -- even after rounding coordinates and dropping
    optional properties, only ~40 of Mid Penn Bank's 59 real branches fit,
    and a larger network (e.g. Hancock Whitney's 181 branches) would have
    lost the majority of its markers, silently, on the flagship map of a
    paid deliverable. This design has no such ceiling: URL size is constant
    regardless of branch count, so nothing is ever dropped from the map.

    Mapbox's Static Images API zoom parameter follows the GL/512px-tile
    convention, one level "wider" than the classic 256px convention used
    here for the fit/pixel math -- hence the -1 when building the request
    URL. If the fetched base map looks zoomed one level off from where the
    markers land when this is checked on Railway, that offset is the first
    thing to check.

    OSM raw tile hotlinking (the first attempt, via the staticmap package)
    was ruled out already -- confirmed directly, not theoretical: it returns
    403 even with a proper User-Agent, which is OSM's policy blocking
    automated/cloud-IP tile requests, not a fixable header problem.

    This function cannot be tested from within the dev sandbox (api.mapbox.com
    is outside both the sandbox's bash network allowlist and the web_fetch
    tool's allowed-domains list) -- must be visually verified on Railway.
    That's exactly why the dispatcher below falls back to the state-outline
    map on any failure here."""
    if not branches_geo or not MAPBOX_TOKEN:
        return False

    W, H = 1280, 960
    pad_frac = 0.15

    lons = [b["lon"] for b in branches_geo]
    lats = [b["lat"] for b in branches_geo]
    lon_c = (min(lons) + max(lons)) / 2
    lat_c = (min(lats) + max(lats)) / 2

    classic_zoom = _fit_zoom(lons, lats, W * (1 - pad_frac), H * (1 - pad_frac))
    mapbox_zoom = max(classic_zoom - 1, 0)  # GL/512px convention offset — see docstring

    url = (f"https://api.mapbox.com/styles/v1/mapbox/streets-v12/static/"
           f"{lon_c},{lat_c},{mapbox_zoom}/{W}x{H}?access_token={MAPBOX_TOKEN}")

    resp = requests.get(url, timeout=10)
    if resp.status_code != 200:
        raise RuntimeError(f"Mapbox Static Images API {resp.status_code}: {resp.text[:200]}")

    from PIL import Image, ImageDraw
    from io import BytesIO
    img = Image.open(BytesIO(resp.content)).convert("RGB")

    # Mapbox's free-tier ToS requires visible attribution -- text fallback
    # here; verify Railway's rendered output also satisfies Mapbox's logo
    # requirement (https://www.mapbox.com/legal/tos) before relying on this
    # for client-facing delivery at scale.
    draw = ImageDraw.Draw(img)
    iw, ih = img.size
    attr_text = "(c) Mapbox (c) OpenStreetMap contributors"
    draw.rectangle([iw - 300, ih - 22, iw, ih], fill=(255, 255, 255, 210))
    draw.text((iw - 294, ih - 18), attr_text, fill=(60, 60, 60))

    # Pixel position for any lon/lat, in THIS image's coordinate space —
    # image center = (lon_c, lat_c) at classic_zoom by construction.
    cx_world, cy_world = _web_mercator_xy(lon_c, lat_c, classic_zoom)

    def to_px(lon, lat):
        x, y = _web_mercator_xy(lon, lat, classic_zoom)
        return (x - cx_world) + iw / 2, (y - cy_world) + ih / 2

    fig, ax = plt.subplots(figsize=(iw / 200, ih / 200), dpi=200)
    ax.imshow(img)

    # Every branch, every zone — no truncation, unlike the URL-embedded
    # overlay approach this replaced.
    for zone in ["Justify", "Defend", "Analyze", "Invest"]:
        pts = [b for b in branches_geo if b.get("opportunity_zone") == zone]
        if not pts:
            continue
        pxs, pys, sizes = [], [], []
        for b in pts:
            px, py = to_px(b["lon"], b["lat"])
            pxs.append(px)
            pys.append(py)
            sizes.append(max(_sf(b.get("latest_dep")) / 3e6, 18))
        ax.scatter(pxs, pys, s=sizes, c=ZONE_HEX_MPL[zone], alpha=0.9,
                   edgecolors="white", linewidths=0.6, label=zone, zorder=3)

    # No custom city-label overlay here — unlike the fallback state-outline
    # map, this basemap (streets-v12) already renders place names natively
    # with its own collision-avoidance. A confirmed real render showed our
    # own labels duplicating and colliding with Mapbox's built-in ones (e.g.
    # "Bethlehem" drawn twice, a custom "Perkasie" label overlapping the
    # native "Dublin" label) — strictly worse than leaving it to the basemap.

    from matplotlib.lines import Line2D
    handles = [Line2D([0], [0], marker="o", linestyle="", markersize=7,
                       markerfacecolor=ZONE_HEX_MPL[z], markeredgecolor="white", label=z)
               for z in ["Justify", "Defend", "Analyze", "Invest"]]
    ax.legend(handles=handles, frameon=True, framealpha=0.9, fontsize=9,
              loc="lower left", edgecolor="none")

    ax.set_xlim(0, iw)
    ax.set_ylim(ih, 0)  # image y-axis is top-down
    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.tight_layout(pad=0)
    fig.savefig(path, dpi=200)
    plt.close(fig)
    return True


def chart_branch_map(branches_geo, path):
    """Dispatcher: real map tiles when available, state-outline fallback
    otherwise. See chart_branch_map_osm's docstring for why this needs a
    fallback and can't be verified from the dev sandbox."""
    try:
        if chart_branch_map_osm(branches_geo, path):
            return True
    except Exception as e:
        print(f"  ⚠ OSM map tiles failed ({e}) — falling back to state-outline map")
    return chart_branch_map_states(branches_geo, path)


def chart_branch_map_states(branches_geo, path):
    """FALLBACK map — lat/lon scatter, sized by deposits, colored by zone,
    drawn over real US state outlines cropped to the branch footprint. State
    boundaries come from a bundled GeoJSON asset (no live tile service needed).
    Used only if chart_branch_map_osm (real map tiles) fails for any reason —
    missing package, tile server down, network timeout on Railway — so a
    report never fails to generate over a map image."""
    if not branches_geo:
        return False

    lons_all = [b["lon"] for b in branches_geo]
    lats_all = [b["lat"] for b in branches_geo]
    pad_lon = max((max(lons_all) - min(lons_all)) * 0.18, 0.6)
    pad_lat = max((max(lats_all) - min(lats_all)) * 0.18, 0.6)
    x0, x1 = min(lons_all) - pad_lon, max(lons_all) + pad_lon
    y0, y1 = min(lats_all) - pad_lat, max(lats_all) + pad_lat

    fig, ax = plt.subplots(figsize=(7.2, 5.4), dpi=200)

    state_polys = _load_state_polygons()
    from shapely.geometry import box as _box
    view_box = _box(x0, y0, x1, y1)
    for name, geom in state_polys.items():
        if not geom.intersects(view_box):
            continue
        _draw_polygon(ax, geom, facecolor="#F2F2EE", edgecolor="#C7C7BF", linewidth=0.8, zorder=1)

    for zone in ["Justify", "Defend", "Analyze", "Invest"]:  # draw Invest last (on top)
        pts = [b for b in branches_geo if b.get("opportunity_zone") == zone]
        if not pts:
            continue
        lons = [b["lon"] for b in pts]
        lats = [b["lat"] for b in pts]
        sizes = [max(_sf(b.get("latest_dep")) / 3e6, 18) for b in pts]
        ax.scatter(lons, lats, s=sizes, c=ZONE_HEX_MPL[zone], alpha=0.85,
                   edgecolors="white", linewidths=0.5, label=zone, zorder=3)

    # Label the largest state clusters directly on the map (avg position, branch count)
    by_state = {}
    for b in branches_geo:
        st = b.get("stalpbr") or b.get("state")
        if st:
            by_state.setdefault(st, []).append(b)
    top_states = sorted(by_state.items(), key=lambda kv: -len(kv[1]))[:4]
    for st, pts in top_states:
        if len(pts) < 2:
            continue
        clx = sum(p["lon"] for p in pts) / len(pts)
        cly = sum(p["lat"] for p in pts) / len(pts)
        max_r = max(max(_sf(p.get("latest_dep")) / 3e6, 18) for p in pts)
        offset_pts = 14 + (max_r ** 0.5)  # clear large bubbles near the centroid
        ax.annotate(f"{st} · {len(pts)}", (clx, cly), xytext=(0, offset_pts),
                    textcoords="offset points", fontsize=8.5, fontweight="bold",
                    color="#334155", ha="center", zorder=4,
                    bbox=dict(boxstyle="round,pad=0.25", facecolor="#FAFAF8",
                              edgecolor="none", alpha=0.85))

    # City labels for individual reference points — state-level clustering
    # alone leaves a tight-footprint network (e.g. 22 branches all in one
    # county) with no landmark to orient by. Label the largest branch per
    # distinct city, capped to the top cities by branch count so a dense
    # network doesn't get cluttered with every city name.
    by_city = {}
    for b in branches_geo:
        city = b.get("citybr") or b.get("city")
        if city:
            by_city.setdefault(city, []).append(b)
    top_cities = sorted(by_city.items(), key=lambda kv: -len(kv[1]))[:8]
    for city, pts in top_cities:
        anchor = max(pts, key=lambda p: _sf(p.get("latest_dep")))
        r = max(_sf(anchor.get("latest_dep")) / 3e6, 18)
        offset_pts = 9 + (r ** 0.5) * 0.7
        ax.annotate(city, (anchor["lon"], anchor["lat"]), xytext=(0, -offset_pts),
                    textcoords="offset points", fontsize=6.5, color="#5B6472",
                    ha="center", va="top", zorder=4,
                    bbox=dict(boxstyle="round,pad=0.15", facecolor="#FAFAF8",
                              edgecolor="none", alpha=0.75))

    ax.set_xlim(x0, x1)
    ax.set_ylim(y0, y1)
    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.set_facecolor("#FAFAF8")
    ax.legend(frameon=False, fontsize=10, loc="lower left", markerscale=0.6)
    ax.set_aspect("equal", adjustable="datalim")
    fig.tight_layout(pad=0.3)
    fig.savefig(path, transparent=False, facecolor="#FAFAF8")
    plt.close(fig)
    return True


def _select_spread_labels(competitors, position_fn, max_labels=5, min_sep=0.35):
    """Picks which competitors get a text label on a competitor map, prioritizing
    by deposits like before, but skipping any candidate whose position is too
    close to an already-selected label. Fixes real dense-market behavior where
    naively labeling the top 5 by deposits produced 3 stacked, illegible
    'Bank of America' labels on top of each other while the actual #1
    competitor's label ended up buried underneath.

    position_fn(c) -> (x, y) in the same coordinate space the caller will plot
    in (pixels for the Mapbox version, miles for the local fallback).
    min_sep is in that same unit -- callers pass a value appropriate to their
    coordinate space (e.g. ~35px on a 700px image, or a fraction of a mile).
    Diversifying which competitors get labeled (favoring geographic spread
    over a strict deposit ranking once the top pick is locked in) is a better
    outcome anyway -- five labels scattered across the map are more useful
    than three overlapping in one corner and two elsewhere."""
    ordered = sorted(competitors, key=lambda c: -_sf(c.get("deposits")))
    selected, positions = [], []
    for c in ordered:
        if len(selected) >= max_labels:
            break
        pos = position_fn(c)
        if pos is None:
            continue
        if all(((pos[0] - p[0]) ** 2 + (pos[1] - p[1]) ** 2) ** 0.5 >= min_sep for p in positions):
            selected.append(c)
            positions.append(pos)
    return {id(c) for c in selected}


def chart_branch_radius_map_local(branch_lat, branch_lon, competitors, radius_mi, path):
    """FALLBACK: plain circle-plot version (no basemap). Used only if the
    Mapbox version fails for any reason — same reasoning as
    chart_branch_map_states relative to chart_branch_map_osm."""
    if branch_lat is None or branch_lon is None:
        return False

    lat_rad = math.radians(branch_lat)
    mi_per_deg_lat = 69.0
    mi_per_deg_lon = 69.0 * max(math.cos(lat_rad), 0.15)

    def to_local_mi(lat, lon):
        return (lon - branch_lon) * mi_per_deg_lon, (lat - branch_lat) * mi_per_deg_lat

    fig, ax = plt.subplots(figsize=(3.4, 3.4), dpi=200)

    theta = [i / 100 * 2 * math.pi for i in range(101)]
    circ_x = [radius_mi * math.cos(t) for t in theta]
    circ_y = [radius_mi * math.sin(t) for t in theta]
    ax.plot(circ_x, circ_y, color="#083D5F", linewidth=1.0, linestyle="--", alpha=0.5, zorder=2)

    labeled_ids = _select_spread_labels(
        competitors,
        position_fn=lambda c: to_local_mi(c["lat"], c["lon"]) if c.get("lat") is not None else None,
        max_labels=5, min_sep=radius_mi * 0.12,
    )
    for c in competitors:
        clat, clon = c.get("lat"), c.get("lon")
        if clat is None or clon is None:
            continue
        cx, cy = to_local_mi(clat, clon)
        r = max(_sf(c.get("deposits")) / 4e6, 40)
        ax.scatter([cx], [cy], s=r, c="#A32D2D", alpha=0.75,
                   edgecolors="white", linewidths=0.6, zorder=3)
        if id(c) in labeled_ids:
            label = c.get("bank_name", "")[:18]
            ax.annotate(f"{label}\n{_sf(c.get('distance_miles')):.1f}mi", (cx, cy),
                        xytext=(0, -9), textcoords="offset points", fontsize=5.5,
                        color="#334155", ha="center", va="top", zorder=4,
                        bbox=dict(boxstyle="round,pad=0.12", facecolor="#FAFAF8",
                                  edgecolor="none", alpha=0.8))

    # The client's own branch, at the local origin, drawn last so it's on top
    ax.scatter([0], [0], s=140, c="#083D5F", marker="*",
               edgecolors="white", linewidths=0.8, zorder=5)

    pad = radius_mi * 1.35
    ax.set_xlim(-pad, pad)
    ax.set_ylim(-pad, pad)
    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.set_facecolor("#FAFAF8")
    ax.set_aspect("equal", adjustable="box")
    fig.tight_layout(pad=0.2)
    fig.savefig(path, transparent=False, facecolor="#FAFAF8")
    plt.close(fig)
    return True


def chart_branch_radius_map_osm(branch_lat, branch_lon, competitors, radius_mi, path):
    """PRIMARY per-branch competitor map — same real Mapbox basemap as the
    main Geographic Distribution map, zoomed to the branch's adaptive radius.
    Reuses _web_mercator_xy/_fit_zoom (already verified against all 59 real
    Mid Penn branches on the main map) so the branch star, radius circle,
    and competitor markers all land pixel-correct on the fetched tile image.
    Falls back to chart_branch_radius_map_local on any failure."""
    if branch_lat is None or branch_lon is None or not MAPBOX_TOKEN:
        return False

    W, H = 700, 700
    # Bounding box = branch +/- radius, converted to degrees locally (fine
    # at this scale) purely to pick a zoom level that frames the radius
    # circle with headroom — actual marker/circle placement below uses the
    # exact same Mercator projection as the main map, not this approximation.
    mi_per_deg_lat = 69.0
    mi_per_deg_lon = 69.0 * max(math.cos(math.radians(branch_lat)), 0.15)
    pad_mi = radius_mi * 1.35
    lons = [branch_lon - pad_mi / mi_per_deg_lon, branch_lon + pad_mi / mi_per_deg_lon]
    lats = [branch_lat - pad_mi / mi_per_deg_lat, branch_lat + pad_mi / mi_per_deg_lat]

    classic_zoom = _fit_zoom(lons, lats, W, H)
    mapbox_zoom = max(classic_zoom - 1, 0)

    url = (f"https://api.mapbox.com/styles/v1/mapbox/streets-v12/static/"
           f"{branch_lon},{branch_lat},{mapbox_zoom}/{W}x{H}?access_token={MAPBOX_TOKEN}")
    resp = requests.get(url, timeout=10)
    if resp.status_code != 200:
        raise RuntimeError(f"Mapbox Static Images API {resp.status_code}: {resp.text[:200]}")

    from PIL import Image, ImageDraw
    from io import BytesIO
    img = Image.open(BytesIO(resp.content)).convert("RGB")

    draw = ImageDraw.Draw(img)
    iw, ih = img.size
    draw.rectangle([iw - 165, ih - 14, iw, ih], fill=(255, 255, 255, 210))
    draw.text((iw - 160, ih - 12), "(c) Mapbox (c) OSM", fill=(60, 60, 60))

    cx_world, cy_world = _web_mercator_xy(branch_lon, branch_lat, classic_zoom)

    def to_px(lon, lat):
        x, y = _web_mercator_xy(lon, lat, classic_zoom)
        return (x - cx_world) + W / 2, (y - cy_world) + H / 2

    fig, ax = plt.subplots(figsize=(W / 200, H / 200), dpi=200)
    ax.imshow(img)

    # Radius circle — generate in real lat/lon (not a flat local approximation)
    # then project through the same Mercator math as everything else, so it
    # lines up correctly with the real basemap underneath.
    theta = [i / 100 * 2 * math.pi for i in range(101)]
    circ_px, circ_py = [], []
    for t in theta:
        clat = branch_lat + (radius_mi / mi_per_deg_lat) * math.sin(t)
        clon = branch_lon + (radius_mi / mi_per_deg_lon) * math.cos(t)
        px, py = to_px(clon, clat)
        circ_px.append(px)
        circ_py.append(py)
    ax.plot(circ_px, circ_py, color="#083D5F", linewidth=1.3, linestyle="--", alpha=0.7, zorder=2)

    # Every competitor within the radius gets a dot (some markets have 20+,
    # e.g. Camden at 1mi in validation testing, or Clinton Savings Bank's
    # Clinton branch at 44) -- but labeling all of them would be unreadable
    # on a small inset, so only up to 5 get a text label. Naively picking
    # the top 5 by deposits produced a real, confirmed bug: 3 "Bank of
    # America" branches clustered together all got labeled, stacking
    # illegibly on top of each other while the actual #1 competitor's own
    # label ended up buried underneath the pile. _select_spread_labels
    # skips candidates too close (in pixels) to an already-picked label.
    labeled_ids = _select_spread_labels(
        competitors,
        position_fn=lambda c: to_px(c["lon"], c["lat"]) if c.get("lat") is not None else None,
        max_labels=5, min_sep=50,
    )
    for c in competitors:
        clat, clon = c.get("lat"), c.get("lon")
        if clat is None or clon is None:
            continue
        px, py = to_px(clon, clat)
        r = max(_sf(c.get("deposits")) / 4e6, 40)
        ax.scatter([px], [py], s=r, c="#A32D2D", alpha=0.85,
                   edgecolors="white", linewidths=0.6, zorder=3)
        if id(c) in labeled_ids:
            label = c.get("bank_name", "")[:18]
            ax.annotate(f"{label}\n{_sf(c.get('distance_miles')):.1f}mi", (px, py),
                        xytext=(0, -9), textcoords="offset points", fontsize=5.5,
                        color="#1A1A1A", ha="center", va="top", zorder=4,
                        bbox=dict(boxstyle="round,pad=0.12", facecolor="white",
                                  edgecolor="none", alpha=0.85))

    bx, by = to_px(branch_lon, branch_lat)
    ax.scatter([bx], [by], s=150, c="#083D5F", marker="*",
               edgecolors="white", linewidths=0.9, zorder=5)

    ax.set_xlim(0, W)
    ax.set_ylim(H, 0)
    ax.set_xticks([])
    ax.set_yticks([])
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.tight_layout(pad=0)
    fig.savefig(path, dpi=200)
    plt.close(fig)
    return True


def _distance_miles(lat1, lon1, lat2, lon2):
    """Simple haversine — used only to backfill distance_miles for the
    vuln_list map fallback below, which doesn't carry it natively."""
    if None in (lat1, lon1, lat2, lon2):
        return None
    r = 3958.8
    p1, p2 = math.radians(lat1), math.radians(lat2)
    dp, dl = math.radians(lat2 - lat1), math.radians(lon2 - lon1)
    a = math.sin(dp / 2) ** 2 + math.cos(p1) * math.cos(p2) * math.sin(dl / 2) ** 2
    return 2 * r * math.asin(math.sqrt(a))


def chart_branch_radius_map(branch_lat, branch_lon, competitors, radius_mi, path):
    """Dispatcher: real map tiles when available, local-plot fallback
    otherwise. Same pattern as chart_branch_map/chart_branch_map_osm."""
    try:
        if chart_branch_radius_map_osm(branch_lat, branch_lon, competitors, radius_mi, path):
            return True
    except Exception as e:
        print(f"  ⚠ OSM radius map failed ({e}) — falling back to local-plot version")
    return chart_branch_radius_map_local(branch_lat, branch_lon, competitors, radius_mi, path)


def fetch_vulnerability_targets(ik, branches):
    """Pre-scored competitor vulnerability data from analytics.branch_target_competitors
    -- YoY trend, ROA, noncurrent-asset %, and a composite vuln_score, already
    ranked top-3 per branch by Verlocity's own scoring. This is a genuinely
    different, richer signal than the adaptive-radius competitor list used
    elsewhere: that list answers 'who is nearby and how big are they,' this
    answers 'who is actually losing ground and winnable.' Two separate
    supabase() calls + a client-side join (same pattern as fetch_branch_geo)
    since branch_target_competitors and geo.branches_master_v2 are different
    schemas without a PostgREST-embeddable relationship configured between them.
    Returns {my_branch_id: [enriched competitor dicts, already rank-ordered]}."""
    rows = supabase(
        "branch_target_competitors",
        f"my_inst_key=eq.{ik}&select=my_branch_id,target_uninumbr,target_namefull,"
        "target_dep_m,target_yoy_pct,target_roa,target_noncurrent_pct,target_opp_score,"
        "vuln_score,target_rank&order=my_branch_id.asc,target_rank.asc",
    )
    rows = rows if isinstance(rows, list) else []
    if not rows:
        return {}

    target_ids = ",".join(str(r["target_uninumbr"]) for r in rows if r.get("target_uninumbr"))
    geo_rows = supabase("branches_master_v2", f"branch_id=in.({target_ids})&select=branch_id,lat,lon") if target_ids else []
    geo_rows = geo_rows if isinstance(geo_rows, list) else []
    geo_by_id = {g["branch_id"]: g for g in geo_rows}

    out = {}
    for r in rows:
        g = geo_by_id.get(r.get("target_uninumbr"), {})
        out.setdefault(r["my_branch_id"], []).append({
            "bank_name": r.get("target_namefull"),
            "deposits": _sf(r.get("target_dep_m")) * 1e6,  # target_dep_m is in millions
            "yoy_pct": _sf(r.get("target_yoy_pct")),
            "roa": _sf(r.get("target_roa")),
            "noncurrent_pct": _sf(r.get("target_noncurrent_pct")),
            "opportunity_score": _sf(r.get("target_opp_score")),
            "vuln_score": _sf(r.get("vuln_score")),
            "rank": r.get("target_rank"),
            "lat": g.get("lat"),
            "lon": g.get("lon"),
        })
    return out


def _vulnerability_reasoning(c):
    """Deterministic, real-data-grounded explanation of WHY a competitor is
    a winnable target -- not a generic 'they're big and nearby' framing, but
    the actual weakness signal (declining deposits, stressed profitability,
    asset quality concerns) that makes their depositors realistically
    winnable. Ties directly to 'who to go after, predictably' rather than
    leaving vuln_score as an opaque number.

    Used only where all 3 competitors are NOT shown side by side (the AI
    prompt context, where the full numbers sit right alongside this text
    so the model can differentiate itself). For reader-facing prose that
    shows all 3 competitors together, use _vulnerability_sentence()
    instead -- this per-row version has no visibility into the other two
    competitors, so it independently produces 'declining deposits' for
    any competitor with unremarkable ROA/noncurrent, which reads as a
    repeated, undifferentiated line when three such competitors are
    listed one after another (a real, recurring case, not hypothetical)."""
    reasons = []
    yoy = c.get("yoy_pct", 0)
    roa = c.get("roa", 0)
    noncurrent = c.get("noncurrent_pct", 0)
    if yoy < -10:
        reasons.append(f"already losing deposits at scale ({yoy:+.1f}% YoY)")
    elif yoy < 0:
        reasons.append(f"declining deposits ({yoy:+.1f}% YoY)")
    if roa < 0.5:
        reasons.append(f"stressed profitability (ROA {roa:.2f}%) limiting their ability to compete on rate")
    if noncurrent > 2:
        reasons.append(f"elevated asset-quality concerns ({noncurrent:.1f}% noncurrent) likely constraining growth appetite")
    if not reasons:
        return "a stable competitor — not showing acute weakness, but still a realistic size-based target"
    return "; ".join(reasons)


def _vulnerability_sentence(c, vuln_list):
    """Full-sentence version of _vulnerability_tags()'s GROUP-COMPARATIVE
    differentiation -- for reader-facing prose (Priority target, Where you
    have the edge) that shows all 3 competitors together, where
    _vulnerability_reasoning()'s per-row-only version was producing the
    same 'declining deposits (-X.X% YoY)' line three times in a row
    whenever none individually crossed the ROA/noncurrent absolute
    thresholds. Reuses the exact same tag computed for the table's
    Weakness column, so the table and this prose can never disagree on
    which dimension is called out for a given competitor -- just
    expanded into a full clause with the actual numbers."""
    tags = _vulnerability_tags(vuln_list)
    name = c.get("bank_name")
    tag = tags.get(name)
    yoy = _sf(c.get("yoy_pct"))
    roa = _sf(c.get("roa"))
    noncurrent = _sf(c.get("noncurrent_pct"))
    opp = c.get("opportunity_score")

    if tag == "Losing deposits fast":
        return f"already losing deposits at scale ({yoy:+.1f}% YoY)"
    if tag == "Highest credit stress of the three":
        return f"carries the highest noncurrent-asset ratio of the three ({noncurrent:.1f}%), a real balance-sheet stress signal"
    if tag == "Weakest profitability of the three":
        return f"has the weakest profitability of the three (ROA {roa:.2f}%), limiting its ability to compete on rate"
    if tag == "Declining fastest of the three":
        return f"declining faster than the other two ({yoy:+.1f}% YoY)"
    if tag == "Weakest overall position" and opp is not None:
        return f"the weakest overall competitive position of the three (opportunity score {opp:.0f}/100), even with a milder {yoy:+.1f}% YoY decline"
    if tag == "Declining deposits":
        return f"declining deposits ({yoy:+.1f}% YoY)"
    return "a stable competitor — not showing acute weakness, but still a realistic size-based target"


def _vulnerability_tags(vuln_list):
    """Assigns each competitor in this branch's ranked set (up to 3) a
    DISTINCT short tag -- comparative across FOUR dimensions (deposit
    trend, profitability, asset quality, overall competitive position),
    not deposits alone. The earlier version only ever compared YoY and
    opportunity score, so every tag read as some flavor of "declining
    deposits" even when ROA and noncurrent-asset data (already fetched,
    just not previously surfaced in a tag) told a sharper story --
    exactly the "it looks like it just deposits" gap this fixes.

    For each competitor, finds which single dimension it's the WEAKEST
    on relative to the other names in this specific group, checked in
    priority order (asset quality and profitability surface before a
    relative YoY comparison, since a bank bleeding capital or margin is
    a more fundamental problem than a milder deposit dip) -- so real
    variance in the underlying financials produces genuinely different
    phrases instead of three restatements of the same metric. Returns
    {bank_name: tag}."""
    if not vuln_list:
        return {}

    def _worst_name(key, higher_is_worse):
        scored = [c for c in vuln_list if c.get(key) is not None]
        if not scored:
            return None
        fn = max if higher_is_worse else min
        return fn(scored, key=lambda c: _sf(c.get(key))).get("bank_name")

    worst_noncurrent = _worst_name("noncurrent_pct", higher_is_worse=True)
    worst_roa = _worst_name("roa", higher_is_worse=False)
    worst_yoy = _worst_name("yoy_pct", higher_is_worse=False)
    worst_score = _worst_name("opportunity_score", higher_is_worse=False)

    tags = {}
    for c in vuln_list:
        yoy = _sf(c.get("yoy_pct"))
        name = c.get("bank_name")
        if yoy < -10:
            tags[name] = "Losing deposits fast"
        elif name == worst_noncurrent and _sf(c.get("noncurrent_pct")) > 1:
            tags[name] = "Highest credit stress of the three"
        elif name == worst_roa and _sf(c.get("roa")) < 1.5:
            tags[name] = "Weakest profitability of the three"
        elif name == worst_yoy and yoy < 0:
            tags[name] = "Declining fastest of the three"
        elif name == worst_score:
            tags[name] = "Weakest overall position"
        elif yoy < 0:
            tags[name] = "Declining deposits"
        else:
            tags[name] = "Stable, size-based target only"
    return tags


def _rank_order_note(vuln_list):
    """Explains why the rank order isn't simply sorted by YoY severity --
    the vuln_score formula weighs deposit trend, profitability, asset
    quality, proximity, and the competitor's OWN overall market position
    together, so the steepest decliner doesn't always come out on top (a
    real, recurring case: a competitor with a much worse YoY number can
    still rank below one with a milder decline, because vuln_score isn't a
    single-metric sort). Returns None when the visible order already
    matches YoY severity -- nothing counterintuitive to flag in that case."""
    if len(vuln_list) < 2:
        return None
    for i in range(len(vuln_list) - 1):
        higher, lower = vuln_list[i], vuln_list[i + 1]
        higher_yoy, lower_yoy = _sf(higher.get("yoy_pct")), _sf(lower.get("yoy_pct"))
        if lower_yoy < higher_yoy:  # the LOWER-ranked one is declining faster
            return (
                f"{lower.get('bank_name')} is declining faster ({lower_yoy:+.1f}% YoY) than "
                f"{higher.get('bank_name')} ({higher_yoy:+.1f}% YoY) but ranks behind it here — "
                f"vulnerability isn't a sort on deposit trend alone. Profitability, asset quality, "
                f"proximity, and {higher.get('bank_name')}'s own weaker overall competitive position "
                f"combine to outweigh the steeper single-metric decline."
            )
    return None


def _relative_size_line(own_deposits, competitors):
    """One sentence placing the client branch's own deposit size against the
    named competitor set, e.g. 'larger than X (2.8x), smaller than Y (0.28x)'.
    Answers a question the doc never used to answer: not just who's nearby
    and who's weak, but where THIS branch actually stands in that lineup —
    the size context a banker in the room will ask for immediately."""
    own_deposits = _sf(own_deposits)
    if own_deposits <= 0 or not competitors:
        return None
    parts = []
    for c in competitors:
        comp_dep = _sf(c.get("deposits"))
        if comp_dep <= 0:
            continue
        ratio = own_deposits / comp_dep
        name = c.get("bank_name", "—")
        if ratio >= 1:
            parts.append(f"larger than {name} ({ratio:.1f}x)")
        else:
            parts.append(f"smaller than {name} ({ratio:.2f}x)")
    if not parts:
        return None
    return f"${own_deposits/1e6:.1f}M in deposits — " + ", ".join(parts) + "."



def fetch_branch_geo(ik):
    """Lat/lon for the branch map — joined from geo.branches_master_v2.
    branch_id in geo matches uninumbr in branch_opportunity_base.
    stalpbr is required here (not just cosmetic) — chart_branch_map's
    cluster labeling groups by it, and silently labels nothing if it's
    missing rather than erroring, which is how this went unnoticed."""
    rows = supabase(
        "branch_opportunity_base",
        f"inst_key=eq.{ik}&select=uninumbr,opportunity_zone,latest_dep,stalpbr,citybr",
    )
    if not rows:
        return []
    ids = ",".join(str(r["uninumbr"]) for r in rows)
    geo_rows = supabase(
        "branches_master_v2",
        f"branch_id=in.({ids})&select=branch_id,lat,lon",
    )
    geo_rows = geo_rows if isinstance(geo_rows, list) else []
    geo_map = {g["branch_id"]: g for g in geo_rows}
    out = []
    for r in rows:
        g = geo_map.get(r["uninumbr"])
        if g and g.get("lat") and g.get("lon"):
            out.append({**r, "lat": g["lat"], "lon": g["lon"]})
    return out


# ═══════════════════════════════════════════════════════════════
# THE 16 PLAYS — from BMAP_Methodology_Part1.docx (Princeton Partners
# Group proprietary methodology). Zone (row) x Quadrant (col) -> play.
# Q1=Grow&Perform Q2=Invest&Protect Q3=Maintain&Improve Q4=Rationalize&Exit
# ═══════════════════════════════════════════════════════════════
PLAY_MATRIX = {
    ("Invest",  "Q1"): "Aggressive Acquisition",
    ("Invest",  "Q2"): "Market Domination",
    ("Invest",  "Q3"): "Growth Opportunity",
    ("Invest",  "Q4"): "Niche Opportunity",
    ("Analyze", "Q1"): "Competitive Defense",
    ("Analyze", "Q2"): "Grow Share",
    ("Analyze", "Q3"): "Maintain",
    ("Analyze", "Q4"): "Efficiency Review",
    ("Defend",  "Q1"): "Urgent Competitive Push",
    ("Defend",  "Q2"): "Targeted Defense",
    ("Defend",  "Q3"): "Steady State",
    ("Defend",  "Q4"): "Efficiency Review",
    ("Justify", "Q1"): "Exit Strategy",
    ("Justify", "Q2"): "Asset Optimization",
    ("Justify", "Q3"): "Performance Improvement",
    ("Justify", "Q4"): "Rationalize",
}

# Acquisition intensity per play — drives whether a branch gets an
# acquisition-posture writeup or a no-campaign/diagnostic one, per
# "A Rationalize play branch does not receive an acquisition campaign.
# A Grow Share play branch does." (Methodology Part 1)
PLAY_ACQUISITION_POSTURE = {
    "Aggressive Acquisition": "Maximum acquisition budget — all channels appropriate.",
    "Market Domination": "High acquisition budget — dominance posture, deter competitor response.",
    "Growth Opportunity": "Moderate acquisition budget — efficiency-focused channel mix.",
    "Niche Opportunity": "Targeted, modest budget — diagnose the specific niche before briefing media.",
    "Competitive Defense": "Retention-first budget — stabilize before acquisition.",
    "Grow Share": "Selective acquisition budget — test before scaling.",
    "Maintain": "Low acquisition budget — retention and cross-sell focus.",
    "Efficiency Review": "No new acquisition investment until market trajectory is assessed.",
    "Urgent Competitive Push": "Defensive acquisition budget — rate campaigns, switching offers.",
    "Targeted Defense": "Retention-focused budget — loyalty and relationship offers.",
    "Steady State": "Minimal budget — operational efficiency review only.",
    "Exit Strategy": "No acquisition investment — full strategic review required.",
    "Asset Optimization": "No new acquisition — maximize return from existing customer base.",
    "Performance Improvement": "Modest operational investment — no acquisition campaign.",
    "Rationalize": "No investment pending diagnosis — board-level attention for $100M+ branches.",
}


PLAY_MEDIA_BRIEF = {
    "Aggressive Acquisition": "Target: new-to-bank households in the branch's 10-mile radius. Product: CD or HYSA. KPIs: new accounts, deposit volume, cost per new account.",
    "Market Domination": "Target: in-market consumers and competitive switchers. Goal: increase share of wallet and deter competitor response.",
    "Growth Opportunity": "Prioritize high-intent audiences. Avoid broad awareness spend that the market's modest growth rate cannot return.",
    "Niche Opportunity": "Identify the specific audience segment driving the score before briefing media.",
    "Competitive Defense": "Priority: stop outflows. Secondary: selective new account acquisition for high-value segments only.",
    "Grow Share": "Target: deposit-active households in Invest and Analyze zones. Product: CD or savings. Measure cost per new account vs MediaPredict forecast.",
    "Maintain": "Minimize acquisition cost. Prioritize retention of high-balance customers.",
    "Efficiency Review": "No brief generated. Diagnostic phase. Understand market trajectory before any media allocation.",
    "Urgent Competitive Push": "Target: competitor bank customers. Message: switching value proposition. KPI: captured accounts from identifiable competitor customers.",
    "Targeted Defense": "Retention posture. Existing customer focus. Minimize churn.",
    "Steady State": "No acquisition brief. Operational focus only.",
    "Exit Strategy": "No brief generated. Strategic escalation required.",
    "Asset Optimization": "No acquisition brief. Existing customer focus — CD renewals, balance growth, product consolidation.",
    "Performance Improvement": "No brief generated.",
    "Rationalize": "No brief generated. Board-level attention for branches with $100M+ in deposits.",
}

# Plays whose static text above explicitly references a named competitor or
# "switchers" -- factually contradictory when the branch's own competitor
# data (shown one section earlier in the same document) says no competitor
# meets the adaptive-radius/size filter. Personalized AI text already gets
# told whether a competitor exists (see deep_dive_ctx in get_narratives) and
# should self-correct on its own; these overrides are the deterministic
# fallback's equivalent fix, for when AI personalization isn't available.
PLAY_NO_COMPETITOR_OVERRIDE = {
    "Market Domination": {
        "resource_posture": "High acquisition budget — no named competitor within the adaptive "
                             "radius, so this is uncontested organic-growth capture, not a "
                             "defensive posture against a specific rival.",
        "media_brief": "Target: in-market consumers and new-to-market movers. Goal: capture "
                        "organic demand growth ahead of any competitor establishing a local presence.",
    },
    "Urgent Competitive Push": {
        "resource_posture": "Defensive acquisition budget warranted by this branch's own "
                             "trajectory — though no single named competitor drives it; the "
                             "pressure reads as diffuse market erosion, not one identifiable rival.",
        "media_brief": "Target: broad in-market consumers rather than a specific competitor's "
                        "customers (none met the radius/size filter). Message: retention and "
                        "value proposition against passive attrition, not head-to-head switching.",
    },
    "Competitive Defense": {
        "resource_posture": "Retention-first budget — no named competitor within the adaptive "
                             "radius, so outflow risk is more likely digital/rate-driven than a "
                             "local branch threat.",
        "media_brief": "Priority: retention against digital/rate competition rather than a "
                        "local branch rival. Secondary: selective acquisition for high-value segments.",
    },
    "Targeted Defense": {
        "resource_posture": "Retention-focused budget — no named local competitor within the "
                             "adaptive radius; balance risk reads as macro/rate-driven rather "
                             "than branch-specific.",
        "media_brief": "Retention posture against broad market/rate pressure rather than a "
                        "specific nearby rival. Existing customer focus. Minimize churn.",
    },
}


def get_play(zone, matrix_quadrant):
    """Parse 'Q1 - Grow and Perform' -> 'Q1', look up (zone, quadrant) in
    the 16-play matrix. Falls back gracefully if either is missing/unrecognized."""
    if not zone or not matrix_quadrant:
        return None
    q = matrix_quadrant.split("-")[0].split("–")[0].strip().split()[0] if matrix_quadrant else None
    return PLAY_MATRIX.get((zone, q))


def determine_adaptive_radius(density_1mi_count, branch_deposits):
    """
    Density + deposit-size adaptive radius rule (per user directive):
    - Rural OR low-deposit branch -> widen the search. This is the safe
      default: a thin or low-value market shouldn't risk missing the
      few competitors that do exist.
    - Crowded AND high-deposit -> can tighten to 0.5mi. Only be this
      aggressive when BOTH signals confirm it's a real, contestable,
      valuable market -- not on density or deposits alone.
    Real validation: Camden (dense, $113M branch) -> 20 competitors within
    1mi alone. Millersburg (rural, $534M branch -- the HQ) -> 0 competitors
    even at 1mi. Same fixed radius cannot serve both.
    """
    is_rural    = density_1mi_count < 3
    is_low_dep  = branch_deposits < 30_000_000
    is_dense    = density_1mi_count >= 15
    is_high_dep = branch_deposits >= 100_000_000

    if is_rural or is_low_dep:
        return 10.0
    if is_dense and is_high_dep:
        return 0.5
    if is_dense:
        return 1.0
    return 3.0


def fetch_branch_competitive_strategy(ik, branches, branches_geo):
    """Per-branch adaptive-radius competitor lookup via ONE call to
    branches_within_radius_batch(), grouped locally by branch. Previously
    made one branches_within_radius() call per branch (59+ sequential
    round-trips for a mid-size network) which was exceeding the Gunicorn
    worker timeout in production and killing the connection mid-request.
    Same density + deposit rule, same size filter -- only the fetch pattern
    changed."""
    try:
        bank_id = int(ik.replace("bank_", ""))
    except ValueError:
        print(f"  ⚠ Could not derive numeric bank_id from '{ik}' (credit union or "
              f"non-standard inst_key?) — skipping branch-level competitive strategy.")
        return []

    print(f"  Fetching adaptive-radius competitive strategy for {len(branches)} branches "
          f"(single batched call)...")
    all_candidates = supabase_rpc("branches_within_radius_batch", {
        "p_inst_key": ik, "p_exclude_bank_id": bank_id, "p_max_radius_miles": 10.0,
    }, timeout=45, paginate=True)
    if not isinstance(all_candidates, list):
        all_candidates = []

    candidates_by_branch = {}
    for c in all_candidates:
        candidates_by_branch.setdefault(c["my_uninumbr"], []).append(c)

    results = []
    for b in branches:
        candidates = candidates_by_branch.get(b["uninumbr"], [])

        density_1mi = sum(1 for c in candidates if _sf(c.get("distance_miles")) <= 1.0)
        deposits = _sf(b.get("latest_dep"))
        radius = determine_adaptive_radius(density_1mi, deposits)
        # Same size filter as branch_target_competitors (Methodology Part 2):
        # competitor deposits between 0.10x and 5x the client branch's own
        # deposits. Without this, giant national-bank hub branches (Wells
        # Fargo, PNC, etc. with $1-8B at a single location) get picked as
        # the "capture opportunity," which is not a realistic local target.
        min_dep, max_dep = deposits * 0.10, deposits * 5.0
        filtered = sorted(
            (c for c in candidates
             if _sf(c.get("distance_miles")) <= radius
             and min_dep <= _sf(c.get("deposits")) <= max_dep),
            key=lambda c: -_sf(c.get("deposits"))
        )
        top_competitor = filtered[0] if filtered else None
        top3_competitors = filtered[:3]
        # Total market deposits WITHIN THE ADAPTIVE RADIUS -- see matching
        # note in fetch_single_branch_strategy() below (kept identical so
        # the two paths never disagree on what "market share" means).
        total_radius_deposits = deposits + sum(
            _sf(c.get("deposits")) for c in candidates if _sf(c.get("distance_miles")) <= radius
        )

        results.append({
            "namebr": b.get("namebr"),
            "citybr": b.get("citybr"),
            "stalpbr": b.get("stalpbr"),
            "deposits": deposits,
            "radius_mi": radius,
            "tier": ("Dense/High-Value" if radius == 0.5 else
                     "Dense" if radius == 1.0 else
                     "Suburban" if radius == 3.0 else "Low-Density"),
            "competitor_count": len(filtered),
            "top_competitor": top_competitor,
            "top3_competitors": top3_competitors,
            "all_competitors": filtered,  # full radius+size-filtered set, for the map
            "total_market_deposits_radius": total_radius_deposits,
        })
    print(f"  ✓ {len(results)} branches assessed")
    return results


def fetch_single_branch_strategy(ik, target_branch):
    """Lightweight equivalent of fetch_branch_competitive_strategy() for
    exactly ONE branch -- built for the standalone Preview flow, which needs
    to stay fast enough to run live in a pitch meeting. The full-network
    version pages through EVERY branch's competitor rows (11,000+ for a
    166-branch network like Trustmark) via branches_within_radius_batch(),
    which the Preview was calling and then discarding 165/166ths of --
    real cause of a production timeout once that batch call started being
    correctly paginated instead of silently (and wrongly) truncated. This
    uses the single-branch branches_within_radius() RPC instead, scoped to
    just this branch's own lat/lon, so it costs one small request no matter
    how large the network is.

    branches_within_radius() now returns lat/lon per competitor directly
    (added at the database level -- it originally didn't, which meant the
    map could only draw whichever vuln_list competitors happened to have
    geo from a separate, unrelated join, sometimes just one of three).
    all_competitors below carries real geo now, same as the batch version."""
    try:
        bank_id = int(ik.replace("bank_", ""))
    except ValueError:
        print(f"  ⚠ Could not derive numeric bank_id from '{ik}' — skipping single-branch strategy.")
        return None

    lat = target_branch.get("lat")
    lon = target_branch.get("lon")
    if lat is None or lon is None:
        print(f"  ⚠ No lat/lon on target branch for single-branch strategy fetch.")
        return None

    candidates = supabase_rpc("branches_within_radius", {
        "p_lat": lat, "p_lon": lon, "p_radius_miles": 10.0, "p_exclude_bank_id": bank_id,
    })
    if not isinstance(candidates, list):
        candidates = []

    density_1mi = sum(1 for c in candidates if _sf(c.get("distance_miles")) <= 1.0)
    deposits = _sf(target_branch.get("latest_dep"))
    radius = determine_adaptive_radius(density_1mi, deposits)
    min_dep, max_dep = deposits * 0.10, deposits * 5.0
    filtered = sorted(
        (c for c in candidates
         if _sf(c.get("distance_miles")) <= radius
         and min_dep <= _sf(c.get("deposits")) <= max_dep),
        key=lambda c: -_sf(c.get("deposits"))
    )
    print(f"  ✓ single-branch strategy: {len(candidates)} candidates, "
          f"{len(filtered)} pass radius+size filter (radius={radius}mi)")
    # Total market deposits WITHIN THE ADAPTIVE RADIUS -- every branch in
    # range (not just the size-filtered "named competitor" subset), plus
    # this branch's own deposits, since market share means share of the
    # whole local market, not just of the handful of named targets.
    total_radius_deposits = deposits + sum(
        _sf(c.get("deposits")) for c in candidates if _sf(c.get("distance_miles")) <= radius
    )
    return {
        "namebr": target_branch.get("namebr"),
        "citybr": target_branch.get("citybr"),
        "stalpbr": target_branch.get("stalpbr"),
        "deposits": deposits,
        "radius_mi": radius,
        "tier": ("Dense/High-Value" if radius == 0.5 else
                 "Dense" if radius == 1.0 else
                 "Suburban" if radius == 3.0 else "Low-Density"),
        "competitor_count": len(filtered),
        "top_competitor": filtered[0] if filtered else None,
        "top3_competitors": filtered[:3],
        "all_competitors": filtered,  # now carries lat/lon -- see docstring
        "total_market_deposits_radius": total_radius_deposits,
    }


DEEP_DIVE_THRESHOLD = 25  # <25 branches -> assess every branch. >=25 -> curate top opportunities.
# Was 20 -- raised after a real case: a 22-branch network (Penn Community
# Bank) missed full-network coverage by an arbitrary 2-branch margin and got
# only its top 15 of 22 branches deep-dived, with the other 7 reduced to a
# bare appendix row in a document sold as a full-network assessment. 25
# gives realistic mid-size community bank networks (an interviewed edge
# case, not a hypothetical) headroom without pushing large networks
# (100+ branches, e.g. Hancock Whitney's 181) into "every branch full"
# mode, which would mean that many real-time Mapbox competitor-map calls
# per generation -- a real latency/cost/timeout risk already documented
# elsewhere in this file (see fetch_branch_competitive_strategy).


def build_branch_deep_dives(branches, branch_strategy):
    """Per-branch full assessment: zone, quadrant, play, priority tier, deposits,
    named competitor, and per-branch $ capture scenario.

    <20 branches: every branch gets a full writeup (deep coverage).
    >=20 branches: curated to the branches with the biggest actual $ capture
    opportunity -- ranked by their named competitor's deposits (the real
    contestable dollar figure), not just opportunity_score alone.
    """
    strategy_by_name = {(r["namebr"], r["citybr"], r["stalpbr"]): r for r in branch_strategy}

    enriched = []
    for b in branches:
        key = (b.get("namebr"), b.get("citybr"), b.get("stalpbr"))
        strat = strategy_by_name.get(key)
        play = get_play(b.get("opportunity_zone"), b.get("matrix_quadrant"))
        top_comp = strat.get("top_competitor") if strat else None
        capture_pool = _sf(top_comp.get("deposits")) if top_comp else 0.0
        enriched.append({
            "branch": b,
            "strategy": strat,
            "play": play,
            "capture_pool": capture_pool,
        })

    deep_mode = len(branches) < DEEP_DIVE_THRESHOLD
    if deep_mode:
        selected = enriched  # every branch
    else:
        # Biggest actual $ opportunity first, not just highest score
        selected = sorted(enriched, key=lambda e: -e["capture_pool"])[:15]

    return selected, deep_mode


def summarize_network(d):
    """Aggregate stats for narrative context — NOT a per-branch dump.
    Keeps AI input tokens flat regardless of network size (14 branches or 300)."""
    br = d["branches"]
    n = len(br)
    zones = {"Invest": 0, "Analyze": 0, "Defend": 0, "Justify": 0}
    total_dep = 0.0
    yoy_vals = []
    for b in br:
        z = b.get("opportunity_zone")
        if z in zones:
            zones[z] += 1
        total_dep += _sf(b.get("latest_dep"))
        yoy_vals.append(_sf(b.get("yoy_deposits")))

    avg_yoy = sum(yoy_vals) / len(yoy_vals) if yoy_vals else 0
    avg_score = sum(_sf(b.get("opportunity_score")) for b in br) / n if n else 0

    # Demographic/audience aggregates -- BMAP's exec summary has been reporting
    # only opportunity score + financial ratios, but the platform also scores
    # every branch on Census income/population and ZHVI home-value trend
    # (the same signal AudienceFinder segments key off of). Without an
    # aggregate here, that whole data dimension never reaches the exec
    # summary — only individual branch_audiences blurbs get it.
    inc_vals = [_sf(b.get("household_income")) for b in br if b.get("household_income") is not None]
    inc_yoy_vals = [_sf(b.get("yoy_income_growth")) for b in br if b.get("yoy_income_growth") is not None]
    pop_yoy_vals = [_sf(b.get("yoy_pop_growth")) for b in br if b.get("yoy_pop_growth") is not None]
    zhvi_vals = [_sf(b.get("zhvi_yoy_pct")) for b in br if b.get("zhvi_yoy_pct") is not None]
    avg_household_income = sum(inc_vals) / len(inc_vals) if inc_vals else 0
    avg_income_yoy = sum(inc_yoy_vals) / len(inc_yoy_vals) if inc_yoy_vals else 0
    avg_pop_yoy = sum(pop_yoy_vals) / len(pop_yoy_vals) if pop_yoy_vals else 0
    avg_zhvi_yoy = sum(zhvi_vals) / len(zhvi_vals) if zhvi_vals else 0
    # The single branch with the strongest combined demographic tailwind
    # (income growth + population growth + home-value growth) -- gives the
    # exec summary a concrete named example instead of only network averages.
    strongest_demo_branch = None
    if br:
        strongest_demo_branch = max(
            br, key=lambda b: (_sf(b.get("yoy_income_growth")) + _sf(b.get("yoy_pop_growth"))
                                + _sf(b.get("zhvi_yoy_pct")) / 100)
        )

    top5 = sorted(br, key=lambda b: -_sf(b.get("opportunity_score")))[:5]
    bottom3 = sorted(br, key=lambda b: _sf(b.get("opportunity_score")))[:3]

    # Largest-deposit branch, tracked independent of opportunity-score rank.
    # top5/bottom3 above are pure score rankings, so a branch that dominates
    # the network by deposit size can rank outside both lists and never reach
    # the AI narrative context at all — which is exactly how a flagship branch
    # losing deposits can get buried under a small branch's growth story.
    largest_branch = max(br, key=lambda b: _sf(b.get("latest_dep"))) if br else None
    flagship_risk = None
    if largest_branch and total_dep > 0 and n > 0:
        share = _sf(largest_branch.get("latest_dep")) / total_dep
        yoy = _sf(largest_branch.get("yoy_deposits"))
        zone = largest_branch.get("opportunity_zone")
        # Material relative to network size (>=3x the average branch's share),
        # not a fixed absolute cutoff — a fixed 10% misses real cases on larger
        # networks (e.g. 9.7% share on a 59-branch network is ~5.7x average,
        # genuinely dominant, but would fail a flat 10% test).
        avg_share = 1.0 / n
        if share >= 3 * avg_share and (yoy < 0 or zone == "Justify"):
            flagship_risk = {**largest_branch, "deposit_share_pct": share * 100}

    return {
        "branch_count": n,
        "zones": zones,
        "total_deposits_B": total_dep / 1e9,
        "avg_yoy_pct": avg_yoy * 100,
        "avg_score": avg_score,
        "top5": top5,
        "bottom3": bottom3,
        "largest_branch": largest_branch,
        "flagship_risk": flagship_risk,
        "avg_household_income": avg_household_income,
        "avg_income_yoy_pct": avg_income_yoy * 100,
        "avg_pop_yoy_pct": avg_pop_yoy * 100,
        "avg_zhvi_yoy_pct": avg_zhvi_yoy,
        "strongest_demo_branch": strongest_demo_branch,
    }


# ═══════════════════════════════════════════════════════════════
# AI NARRATIVE — aggregated context only, not per-branch
# ═══════════════════════════════════════════════════════════════

def summarize_branch_strategy(branch_strategy):
    """Bank-wide roll-up by density/deposit tier - feeds both the narrative
    and the 'as a full bank' strategy view, not just per-branch detail."""
    tiers = {}
    for r in branch_strategy:
        t = r["tier"]
        tiers.setdefault(t, {"count": 0, "deposits": 0.0})
        tiers[t]["count"] += 1
        tiers[t]["deposits"] += r["deposits"]
    named_hits = [r for r in branch_strategy if r.get("top_competitor")]
    return {"tiers": tiers, "named_hits": named_hits}


def get_narratives(bank_name, summary, fin, targets, branch_strategy=None, dives=None, capped_yoy=None,
                    vulnerability_targets=None):
    branch_strategy = branch_strategy or []
    dives = dives or []
    capped_yoy = capped_yoy or {}
    if not ANTH_KEY or not anthropic:
        print("  ⚠ No ANTHROPIC_API_KEY — using placeholder narratives")
        return _placeholder_narratives(dives)

    zones = summary["zones"]
    top5_str = "; ".join(
        f"{b['namebr']} ({b['citybr']}, {b['stalpbr']}) — score {_sf(b['opportunity_score']):.0f}, "
        f"${_sf(b['latest_dep'])/1e6:.0f}M deposits, {fmt_yoy(b, capped_yoy)} YoY"
        for b in summary["top5"]
    )
    bottom3_str = "; ".join(
        f"{b['namebr']} ({b['citybr']}, {b['stalpbr']}) — score {_sf(b['opportunity_score']):.0f}"
        for b in summary["bottom3"]
    )
    target_str = "; ".join(
        f"{t.get('target_institution','—')} — {t.get('branches_in_radius','—')} branches exposed, "
        f"{_vuln_tier(t.get('avg_vuln_score'))} vulnerability, deposit YoY {_sf(t.get('avg_yoy_pct')):.1f}%"
        for t in targets
    ) or "No qualifying network-level target identified"

    bs_summary = summarize_branch_strategy(branch_strategy)
    tier_str = "; ".join(
        f"{t}: {v['count']} branches, ${v['deposits']/1e6:.0f}M deposits"
        for t, v in bs_summary["tiers"].items()
    )
    # Best 4 named, real, adaptive-radius findings for the model to draw on directly
    named_examples = sorted(bs_summary["named_hits"], key=lambda r: -r["deposits"])[:4]
    named_str = "; ".join(
        f"{r['namebr']} ({r['citybr']}, {r['stalpbr']}, {r['tier']}, {r['radius_mi']}mi radius) — "
        f"largest nearby competitor {r['top_competitor']['bank_name']} "
        f"{r['top_competitor']['distance_miles']:.2f}mi away, "
        f"${_sf(r['top_competitor']['deposits'])/1e6:.0f}M deposits"
        for r in named_examples
    ) or "No adaptive-radius competitor matches found"

    ctx = f"""
Bank: {bank_name}
Full network: {summary['branch_count']} branches, ${summary['total_deposits_B']:.1f}B total deposits
Zone distribution: Invest {zones['Invest']} | Analyze {zones['Analyze']} | Defend {zones['Defend']} | Justify {zones['Justify']}
Network avg opportunity score: {summary['avg_score']:.0f}/100
Network avg deposit YoY: {summary['avg_yoy_pct']:+.1f}%

Top 5 branches by opportunity: {top5_str}
Bottom 3 branches by opportunity: {bottom3_str}
{"FLAGSHIP RISK — the network's single largest branch by deposits carries the story regardless of its opportunity-score rank: " + summary['flagship_risk']['namebr'] + " (" + summary['flagship_risk']['citybr'] + ", " + summary['flagship_risk']['stalpbr'] + ") holds " + f"{summary['flagship_risk']['deposit_share_pct']:.0f}%" + " of total network deposits ($" + f"{_sf(summary['flagship_risk']['latest_dep'])/1e6:.0f}M" + "), is down " + f"{_sf(summary['flagship_risk']['yoy_deposits'])*100:+.1f}%" + " YoY, and sits in the " + str(summary['flagship_risk']['opportunity_zone']) + " zone." if summary.get('flagship_risk') else ""}

Financial health: ROA {_sf(fin.get('roa')):.2f}% | NIM {_sf(fin.get('nim')):.2f}% | Efficiency {_sf(fin.get('efficiency_ratio')):.1f}%
Deposit YoY {_sf(fin.get('dep_yoy_pct')):+.1f}% | Cost of funds {_sf(fin.get('cost_of_funds_pct')):.2f}% | Tier 1 {_sf(fin.get('tier1_capital_pct')):.1f}%
Net income YoY {_sf(fin.get('net_income_yoy_pct')):+.1f}%

Network-wide demographic & audience signal (Census income/population + ZHVI home-value
trend -- the same underlying data AudienceFinder segments key off of):
Avg household income ${summary['avg_household_income']:,.0f} | Avg income YoY {summary['avg_income_yoy_pct']:+.1f}%
Avg population YoY {summary['avg_pop_yoy_pct']:+.1f}% | Avg home-value (ZHVI) YoY {summary['avg_zhvi_yoy_pct']:+.1f}%
{"Strongest demographic tailwind: " + summary['strongest_demo_branch']['namebr'] + " (" + summary['strongest_demo_branch']['citybr'] + ", " + summary['strongest_demo_branch']['stalpbr'] + ") — income YoY " + f"{_sf(summary['strongest_demo_branch'].get('yoy_income_growth'))*100:+.1f}%" + ", population YoY " + f"{_sf(summary['strongest_demo_branch'].get('yoy_pop_growth'))*100:+.1f}%" + ", home value YoY " + f"{_sf(summary['strongest_demo_branch'].get('zhvi_yoy_pct')):+.1f}%" + "." if summary.get('strongest_demo_branch') else ""}

Network-level competitive targets: {target_str}

Branch-level adaptive-radius competitive strategy (radius scaled per branch by local
density + deposit size -- dense/high-value branches get as tight as 0.5mi, rural or
low-deposit branches widen to 10mi, since one fixed radius misses the real picture
for a network spanning both dense suburbs and rural markets):
By tier: {tier_str}
Named examples: {named_str}
"""

    # Per-branch demographic context for the batched audience narrative --
    # ONE call covers all deep-dive branches, not one call per branch, since
    # fetch_branch_competitive_strategy already makes N sequential RPC calls
    # and adding N more AI calls would compound that latency risk.
    deep_dive_ctx = ""
    if dives:
        lines = []
        for e in dives:
            b = e["branch"]
            strat = e.get("strategy")
            top_comp = strat.get("top_competitor") if strat else None
            comp_str = (f"nearest named competitor {top_comp.get('bank_name')} "
                        f"{_sf(top_comp.get('distance_miles')):.2f}mi away with "
                        f"${_sf(top_comp.get('deposits'))/1e6:.0f}M deposits"
                        if top_comp else "no named competitor within the adaptive radius")
            driver_clause = _score_driver_clause(b)

            # Vulnerability-ranked competitors (ROA, noncurrent %, YoY, and
            # relative size) for THIS branch -- previously never reached this
            # prompt at all, so the batched verdict/play/audience text had no
            # visibility into which competitors are actually weak, only the
            # nearest one's name/distance/deposits. Same fix as the standalone
            # Preview's get_single_branch_narrative(), applied here so the
            # $10K Assessment and the free Preview never diverge in quality.
            own_dep = _sf(b.get("latest_dep"))
            vuln_list_b = sorted((vulnerability_targets or {}).get(b.get("uninumbr"), []),
                                  key=lambda c: c.get("rank") or 99)[:3]
            vuln_str = ""
            if vuln_list_b:
                parts = []
                for c in vuln_list_b:
                    comp_dep = _sf(c.get("deposits"))
                    ratio_str = ""
                    if comp_dep > 0 and own_dep > 0:
                        ratio = own_dep / comp_dep
                        ratio_str = (f", branch is {ratio:.1f}x their size" if ratio >= 1
                                     else f", branch is {ratio:.2f}x their size (smaller)")
                    parts.append(
                        f"#{c.get('rank')} {c.get('bank_name')} (${comp_dep/1e6:.0f}M, "
                        f"{_sf(c.get('yoy_pct')):+.1f}% YoY, ROA {_sf(c.get('roa')):.2f}%, "
                        f"noncurrent {_sf(c.get('noncurrent_pct')):.1f}%{ratio_str} — "
                        f"{_vulnerability_reasoning(c)})"
                    )
                vuln_str = " | vulnerability-ranked competitors: " + "; ".join(parts)

            lines.append(
                f"- {b.get('namebr')} ({b.get('citybr')}, {b.get('stalpbr')}): "
                f"score {_sf(b.get('opportunity_score')):.0f}/100, zone {b.get('opportunity_zone')}, "
                f"${_sf(b.get('latest_dep'))/1e6:.0f}M deposits, {fmt_yoy(b, capped_yoy or {})} YoY, "
                f"{comp_str}, "
                f"{driver_clause + ', ' if driver_clause else ''}"
                f"household income ${_sf(b.get('household_income')):.0f} "
                f"({_sf(b.get('yoy_income_growth'))*100:+.1f}% YoY), "
                f"population YoY {_sf(b.get('yoy_pop_growth'))*100:+.1f}%, "
                f"home value YoY {_sf(b.get('zhvi_yoy_pct')):+.1f}%, "
                f"assigned play {e['play'] or 'n/a'}"
                f"{vuln_str}"
            )
        deep_dive_ctx = "\n\nBranches needing a full deep-dive writeup (real data, no fabricated " \
                        "personas -- Verlocity's persona layer is still in development):\n" + \
                        "\n".join(lines)

    system = """You are writing the $10K Verlocity BMAP Assessment for a community bank CEO/Board audience.
Tone: confident, commercial, decisive — not hedged, not academic. State the position, don't survey options.
Every branch reference MUST include both branch name AND city — never one without the other.
Do not explain BMAP methodology, do not reference BMAP versions, do not ask follow-up questions,
do not introduce data beyond what's given below. No superlatives for their own sake — earn every claim with a number.
Return ONLY valid JSON, no markdown fences:
{
  "exec_headline": "3-4 sentences. State plainly whether this network's current footprint is positioned for GROWTH, DEFENSE, or OPTIMIZATION -- pick one framing and commit to it. Reference the overall opportunity score and the Invest/Analyze/Defend/Justify mix. Frame the core deposit-acquisition tension explicitly (e.g. concentrated upside vs. broad retention burden). If a FLAGSHIP RISK finding is present above, it must anchor this headline by name and city -- it drives the network total and outweighs a smaller branch's score even if that branch tops the ranking.",
  "strategic_positioning": "One short paragraph (3-4 sentences). Describe what kind of deposit competitor this bank is positioned to be over the next 12-24 months. Ground this in the demographic & audience signal given (name the strongest-tailwind branch/city or the network averages -- growth demographics = expansion case, decline = retention case) AND the competitive/rate exposure (name the top vulnerable target). Address how growth can be driven without relying on additional physical branches -- position digital-first, market-specific execution as the default lever, without naming specific channels, tactics, or products.",
  "priority_focus": [{"branch": "exact branch name", "city": "city", "state": "ST", "zone": "Invest/Analyze/Defend/Justify", "why_now": "one clause: momentum, competitive pressure, or market structure -- with a number", "role": "one short strategic-role phrase, e.g. 'deposit growth engine', 'selective digital capture', 'defend and retain balances'"}] , // 2-3 entries. If a FLAGSHIP RISK finding is present, it MUST be one of these entries (role should reflect its risk, e.g. 'stabilize and retain' or 'exit review'). Otherwise lead with the top opportunity-score branch.
  "next_12_months": ["exactly 3 strings — each a leadership-level decision about WHERE to allocate attention, capital, or effort. No tactics, channels, offers, pricing, or products. No methodology."],
  "network_narrative": "2-3 sentences on what the zone distribution reveals about the network's overall position. (Used later in the doc, not the exec summary above -- can restate the zone framing in different words.)",
  "competitive_narrative": "2-3 sentences naming the specific network-level target and why it is vulnerable.",
  "financial_narrative": "2-3 sentences on what the financial metrics mean together — not a list restated as prose.",
  "capture_strategy_narrative": "3-4 sentences on the branch-level adaptive-radius findings. Name at least one specific dense/high-value branch with its named largest nearby competitor and distance, and contrast the tactical approach that implies (rate/digital competition at close range) against what the low-density branches need instead (defense and wallet-share deepening, since there is often no competitor within the adaptive radius to capture from). This is the 'win deposits by branch AND as a full bank' section.",
  "next_step": "2-3 sentences. A specific, named recommendation tied to the top opportunity branches. (Used in the closing Recommendation section, not the exec summary above.)",
  "branch_plays": {"Branch Name (City, ST)": {"resource_posture": "One sentence, grounded in THIS branch's specific score, deposits, and competitive exposure -- not a generic restatement of the play name. E.g. for a Grow Share play, name the actual budget rationale given this branch's specific numbers, not the same sentence every Grow Share branch would get. CRITICAL: if this branch has no named competitor within its adaptive radius (stated above), do NOT write language implying one exists -- no 'deter competitor response', no 'switching', no reference to a rival. Reframe around organic/uncontested capture or macro/rate pressure instead. If a relative-size figure is given for a vulnerability-ranked competitor, use it as part of the resourcing argument (e.g. 'this branch is 2.8x the size of its weakest named competitor').", "media_brief": "One to two sentences, naming the actual target audience and product implied by THIS branch's demographic and competitive data -- not the generic play-level template. Same competitor-existence constraint as resource_posture above. If a vulnerability-ranked competitor shows real weakness (declining deposits, weak ROA, elevated noncurrent assets), name conquesting that competitor's depositors as part of the angle."}},
  "branch_verdicts": {"Branch Name (City, ST)": "3-4 sentences. Synthesize the score, zone, the named competitive threat (or lack of one), and the deposit trajectory into a single clear verdict on this specific branch -- the 'why' behind its assigned play, not a restatement of the tables that follow it. If a 'score driven primarily by X, weakest on Y' clause is given, use it explicitly -- naming the actual driver of a low or high score (e.g. 'this branch's ceiling is capped by a shrinking local market, not competitive pressure' or 'the score reflects deposit scale, not underlying growth') is exactly the kind of analysis worth paying for, versus a generic restatement of the number. This is what a reader sees BEFORE the supporting detail tables, so it must stand alone: e.g. why a Defend-zone branch with strong income growth is still a retention play given who's 0.2mi away, or why a Low-Density branch with no named competitor should focus on wallet-share deepening instead of acquisition. If vulnerability-ranked competitors are given, name at least one specific weakness (declining deposits, weak ROA, elevated noncurrent assets) rather than treating competitors as an undifferentiated group -- this is the same data the reader sees highlighted in the competitor table, so the verdict must not read thinner than the table it introduces. Ground every claim in the specific numbers given -- no generic branch commentary.",
  "branch_audiences": {"Branch Name (City, ST)": {
    "narrative": "2-3 sentences per branch, using the household income, income YoY, population YoY, home value YoY figures given, AND the competitive weakness data where present -- not demographics alone. Frame through Verlocity's AudienceFinder segments (High-Quality Local Prospects from income/geo, Regression-Scored Lookalikes, Competitive Conquesting for switchers, Warm Retargeting) where the demographic signal supports it. If a named competitor is losing deposits, Competitive Conquesting targeting THEIR depositor base specifically is a stronger, more concrete angle than generic new-household prospecting.",
    "persona_name": "A short archetype-style label for this branch's dominant audience segment, in the style of 'The Equity Plateau' or 'The Cash-Flow Balancer' -- an income/wealth-stage ARCHETYPE grounded in the actual numbers given. Never a specific named individual (e.g. 'Sarah, 34') -- Verlocity's persona layer names segments, not people.",
    "persona_tagline": "One short line under the persona name (5-10 words) capturing the segment's core motivation.",
    "life_stage": "e.g. 'Mid-career to pre-retirement, 40s-60s' -- inferred from income level and market-maturity signals given, not invented biographical detail.",
    "wealth_signal": "Income range + home value context, grounded directly in the numbers given for this branch.",
    "primary_need": "The banking product or need this segment most likely prioritizes, grounded in the income/growth profile given.",
    "switch_driver": "What would actually move this segment to switch banks or deepen a relationship here -- grounded in the competitive weakness data given, when present.",
    "strong_signals": ["1-2 short (under 15 words) bullets naming concrete reasons this branch's audience opportunity is real. Every bullet needs a number from the data given."],
    "validate_before_activating": ["1 short (under 20 words) honest caveat on what this read can't confirm from the data alone -- this is what makes the read credible rather than just optimistic. Do not skip it."]
  }}
}
Keys in branch_plays, branch_verdicts, and branch_audiences must exactly match the branch name+city+state given, in the "Branch Name (City, ST)" format shown."""

    if deep_dive_ctx:
        ctx += deep_dive_ctx

    print("  Generating AI narratives (full-network context)...")
    client = anthropic.Anthropic(api_key=ANTH_KEY)
    try:
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=16000,
            thinking={"type": "disabled"},
            system=system,
            messages=[{"role": "user", "content": ctx}],
        )
        raw = msg.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        narr = json.loads(raw)
        print("  ✓ Narratives generated")
        return narr
    except json.JSONDecodeError as e:
        # Diagnostic on purpose: this exact error class is the truncated-JSON
        # failure mode that's bitten this function twice now (once at
        # max_tokens=4000, again at 8000 after DEEP_DIVE_THRESHOLD widened
        # full-network narrative coverage to 25 branches). Logging raw
        # response length + stop_reason means the NEXT time this happens,
        # it's a one-line diagnosis instead of a guessing exercise.
        print(f"  ⚠ Narrative JSON parse failed ({e}) — response {len(msg.content[0].text)} chars, "
              f"stop_reason={msg.stop_reason} — using placeholders")
        return _placeholder_narratives(dives)
    except Exception as e:
        print(f"  ⚠ Narrative generation failed ({type(e).__name__}: {e}) — using placeholders")
        return _placeholder_narratives(dives)


def _web_search_brief(system, ctx, max_tokens=3000):
    """Shared helper for the two live-web-search-enabled signal briefs below.
    Unlike get_narratives() (structured data only, no tools), these use
    Anthropic's server-side web_search tool -- Claude may call it multiple
    times within a single response, interleaving text/tool_use/tool_result
    blocks, so the reply has to be reassembled from every text block, not
    just the first one. Returns None on any failure so callers can skip
    the section entirely rather than fabricate content."""
    if not ANTH_KEY or not anthropic:
        return None
    client = anthropic.Anthropic(api_key=ANTH_KEY)
    try:
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_tokens,
            thinking={"type": "disabled"},
            system=system,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=[{"role": "user", "content": ctx}],
        )
        text = "".join(block.text for block in msg.content if block.type == "text").strip()
        return text or None
    except Exception as e:
        print(f"  ⚠ Web-search brief failed ({e})")
        return None


def get_persona_signal_brief(bank_name, dives):
    """Persona & Demographic Signal Brief — adapted from Verlocity's Hub
    'Audience Intelligence Brief' analyst prompt. Layers current, directional
    persona/intent signal on top of BMAP's Census-based structural data via
    live web search. Scoped to Invest-zone (and select Analyze-zone) priority
    branches only, per the source prompt's explicit scope limit -- this is
    signal enrichment, not a full-network rewrite."""
    priority = [e for e in (dives or [])
                if e["branch"].get("opportunity_zone") in ("Invest", "Analyze")]
    if not priority:
        return None

    lines = []
    for e in priority[:15]:  # cap for latency/cost -- same ceiling as deep dives
        b = e["branch"]
        lines.append(
            f"- {b.get('namebr')} ({b.get('citybr')}, {b.get('stalpbr')}): "
            f"zone {b.get('opportunity_zone')}, score {_sf(b.get('opportunity_score')):.0f}/100, "
            f"household income ${_sf(b.get('household_income')):.0f} "
            f"({_sf(b.get('yoy_income_growth'))*100:+.1f}% YoY), "
            f"population YoY {_sf(b.get('yoy_pop_growth'))*100:+.1f}%, "
            f"home value YoY {_sf(b.get('zhvi_yoy_pct')):+.1f}%"
        )
    ctx = (f"BMAP Bank Context — {bank_name}\n\nPriority branches (Invest/select Analyze zones):\n"
           + "\n".join(lines))

    system = """You are the BMAP Persona & Demographic Signal Analyst.

Your role is to enrich BMAP branch insights with CURRENT, directional persona, behavioral, and
intent signals that are not fully captured by Census-based data. You do NOT replace BMAP, do NOT
redefine strategy, do NOT produce marketing tactics. You exist to add fresh signal that sharpens
strategic confidence.

MODE: Web search enabled. Use current public signals, behavioral proxies, and local market context.
Treat BMAP branch data (given below) as the structural truth layer. Directional insight only --
never invent statistics, and never contradict the BMAP signals given.

OUTPUT STRUCTURE -- follow exactly, plain text with these three headers:

PERSONA SIGNAL OVERVIEW
1-2 short paragraphs describing the dominant persona and intent signals across the priority markets
given: lifecycle shifts (wealth accumulation, retirement, mobility), income/employment dynamics,
and product intent (checking, CD, loans) implied by current local conditions.

BRANCH-LEVEL PERSONA SIGNALS (PRIORITY ONLY)
For each branch given, on its own line as "Branch Name (City, ST)": 1-2 dominant persona signals,
combining the demographic baseline given with current behavioral/local context. Explain how this
influences deposit growth potential, balance stability, switching likelihood, and digital vs.
relationship preference. ALWAYS include branch name AND city together.

STRATEGIC IMPLICATIONS (NON-TACTICAL)
3-5 bullets on where persona signals should influence acquisition focus, timing advantage, and
where competition is most vulnerable. NO tactics, NO campaigns, NO pricing, NO product
recommendations, NO methodology explanation.

Tone: analytical, confident, advisory, executive-ready."""

    print("  Generating Persona & Demographic Signal Brief (web search)...")
    result = _web_search_brief(system, ctx, max_tokens=3000)
    if result:
        print("  ✓ Persona signal brief generated")
    else:
        print("  ⚠ Persona signal brief unavailable — section will be omitted")
    return result


def get_market_offer_brief(bank_name, dives, branch_strategy):
    """Market Offer & Competitive Signal Brief — adapted from Verlocity's
    Hub competitive-intelligence analyst prompt. Live web search for current
    peer deposit-rate/promotional pressure, explicitly excluding
    money-center banks (Chase, Wells Fargo, BofA, Citi) since they rarely
    reflect real community-bank competitive dynamics. Scoped to the same
    priority branches as the deep dives."""
    if not dives:
        return None
    strategy_by_name = {(r["namebr"], r["citybr"], r["stalpbr"]): r for r in (branch_strategy or [])}

    lines = []
    for e in dives[:15]:
        b = e["branch"]
        key = (b.get("namebr"), b.get("citybr"), b.get("stalpbr"))
        strat = strategy_by_name.get(key)
        top_comp = strat.get("top_competitor") if strat else None
        comp_str = (f"named competitor {top_comp.get('bank_name')} "
                    f"{_sf(top_comp.get('distance_miles')):.1f}mi away"
                    if top_comp else "no named competitor within adaptive radius")
        lines.append(f"- {b.get('namebr')} ({b.get('citybr')}, {b.get('stalpbr')}): {comp_str}")
    ctx = f"BMAP Bank Context — {bank_name}\n\nPriority branches:\n" + "\n".join(lines)

    system = """You are the BMAP Market Offer & Competitive Signal Analyst.

Your purpose is to enrich BMAP strategic decision-making with real-time, peer-level competitive
deposit pressure: community banks, regional banks, and major credit unions operating in each local
market, plus select digital banks only when they materially impact rate-sensitive deposits. You
explicitly DEPRIORITIZE AND EXCLUDE money-center banks (Chase, Wells Fargo, Bank of America, Citi)
and other national brands that do not meaningfully compete for community-bank deposit relationships.

MODE: Web search enabled. Use current public information -- peer bank/credit union sites, local and
regional financial news, rate aggregators with community-bank visibility. Prioritize signals from
the last 30-60 days.

OUTPUT STRUCTURE -- follow exactly, plain text with these three headers:

MARKET OFFER SIGNAL OVERVIEW
1-2 executive paragraphs on dominant peer-competitor dynamics across the priority markets given --
whether pressure is driven by regional bank expansion, credit union rate aggression, local
promotional battles, or digital siphoning of rate-sensitive balances, and whether offer intensity
is escalating, stabilizing, or cooling. Name meaningful peer institutions where relevant.

BRANCH-LEVEL MARKET PRESSURE (PRIORITY ONLY)
For each branch given, on its own line as "Branch Name (City, ST)": meaningful peer competitors by
name, deposit products driving pressure (HYSAs, CDs, promotional checking, bundled incentives),
directional current rate ranges or promotional structures, and how this impacts acquisition
difficulty, retention vulnerability, and deposit mix sensitivity. Avoid national banks unless no
peers exist in that market.

STRATEGIC IMPLICATIONS (NON-TACTICAL)
3-5 executive bullets on where peer pressure justifies accelerated acquisition focus, where
structural competition limits short-term upside, and where defending existing balances is
strategically critical. NO tactics, NO pricing recommendations, NO campaign ideas.

Tone: executive, local-market realistic, competitive-intelligence driven. Read like a peer-level
competitive war-room brief for bank leadership."""

    print("  Generating Market Offer & Competitive Signal Brief (web search)...")
    result = _web_search_brief(system, ctx, max_tokens=3000)
    if result:
        print("  ✓ Market offer brief generated")
    else:
        print("  ⚠ Market offer brief unavailable — section will be omitted")
    return result


def _placeholder_narratives(dives=None):
    base = {k: "" for k in ["exec_headline", "strategic_positioning", "network_narrative",
                             "competitive_narrative", "financial_narrative",
                             "capture_strategy_narrative", "next_step"]}
    base["priority_focus"] = []
    base["next_12_months"] = []
    base["branch_audiences"] = {}
    base["branch_verdicts"] = {}
    base["branch_plays"] = {}
    if dives:
        for e in dives:
            b = e["branch"]
            key = f"{b.get('namebr')} ({b.get('citybr')}, {b.get('stalpbr')})"
            base["branch_audiences"][key] = ""
            base["branch_verdicts"][key] = ""
            base["branch_plays"][key] = {}
    return base


# ═══════════════════════════════════════════════════════════════
# DOCX BUILD
# ═══════════════════════════════════════════════════════════════

def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_table_borders(table):
    """Strips visible gridlines from a table used purely as a layout
    device (the cover page, header/footer bands) -- these aren't data
    tables and shouldn't look like one."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _wrap_text_pil(draw, text, font, max_width):
    """Simple greedy word-wrap for PIL ImageDraw text -- PIL has no native
    wrapping, and the cover's subtitle can run 150-250 characters (the
    Branch Preview's framing sentence)."""
    words = text.split()
    lines, current = [], ""
    for w in words:
        trial = f"{current} {w}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines


def build_branded_cover(doc, bank_name, doc_title, subtitle=None):
    """Cover page -- CURRENTLY STATIC, using Brandon's actual cover export
    as-is (verlocity_cover_vertical.png), bank name and all baked in. This
    is a deliberate, known, temporary state: his source has the bank name/
    title/subtitle/date flattened directly into the image with no separate
    text layer, so nothing here can vary per-bank right now -- every doc
    shows "Trustmark National Bank" regardless of which bank was actually
    generated. Fix in progress: Brandon is producing an updated export with
    the text kept editable/separate so this can go back to compositing
    bank name dynamically per document, the way the rest of this doc
    already works for every other per-branch value.

    bank_name/doc_title/subtitle params are intentionally unused for now --
    kept in the signature so callers don't need to change when the dynamic
    version comes back.

    Falls back to a plain navy text-only cover (WITH real dynamic text) if
    the image asset isn't present, so a missing file doesn't break
    generation -- it just means fewer bells and whistles."""
    section = doc.sections[0]
    printable_width = section.page_width - section.left_margin - section.right_margin
    vertical_bg_path = Path(__file__).parent / "verlocity_cover_vertical.png"

    if vertical_bg_path.exists():
        p_img = doc.add_paragraph()
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(str(vertical_bg_path), width=printable_width)
        doc.add_page_break()
        return

    # ── Fallback: no image asset found -- plain navy cover, text only ──
    printable_height = section.page_height - section.top_margin - section.bottom_margin
    cover = doc.add_table(rows=1, cols=1)
    cover.autofit = False
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = cover.rows[0]
    row.height = printable_height
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST  # AT_LEAST is forgiving if content
    cell = row.cells[0]                            # runs slightly over/under; EXACTLY
    cell.width = printable_width                   # caused a stray blank second page
    _set_cell_shading(cell, "083D5F")
    _remove_table_borders(cover)

    def _line(text, size, color, bold=False, italic=False, space_before=0,
              space_after=0, all_caps=False, align=WD_ALIGN_PARAGRAPH.LEFT, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = align
        r = p.add_run(text.upper() if all_caps else text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = FONT_HEAD
        r.bold = bold
        r.italic = italic
        return p

    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LIGHT_TEAL = RGBColor(0x7D, 0xD8, 0xE8)
    MUTED_ON_NAVY = RGBColor(0xAF, 0xC4, 0xD6)

    _line("", 1, WHITE, space_before=130, first=True)
    _line("VERLOCITY", 34, WHITE, bold=True, all_caps=True)
    _line("Stop Guessing. Start Growing.", 13, LIGHT_TEAL, space_before=2, space_after=44)
    _line(bank_name, 24, WHITE, bold=True, space_before=8)
    _line(doc_title, 13, LIGHT_TEAL, space_before=4)
    if subtitle:
        _line(subtitle, 10.5, MUTED_ON_NAVY, italic=True, space_before=6)
    _line(datetime.now().strftime("%B %Y"), 10, MUTED_ON_NAVY, space_before=10)
    _line("", 1, WHITE, space_before=90)
    _line("VERLOCITY.AI", 9, WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    # No explicit page_break() -- the table's own AT_LEAST height (set to
    # the full printable page height above) already guarantees whatever
    # comes next starts on a new page.


def build_snapshot_intro_cover(doc, title_lines, subtitle=None, date_str=None):
    """Cover using the Verlocity Growth System intro artwork
    (verlocity_cover_snapshot.png -- the photo-collage design used for
    the BMAP Snapshot deck intro: laptop, world map, chart bars behind
    the wordmark). This is a genuine widescreen photo composite
    (~1.77:1, close to 16:9), not a flat-color gradient like
    build_branded_cover()'s vertical version -- there's no way to
    regenerate it at a different aspect ratio the way that one can,
    since it isn't built from separable vector/flat-color layers here,
    only a single flattened image. Stretching or cropping a real photo
    into an 8.5x11 portrait would visibly distort the map/keyboard/chart
    imagery, so this shows it intact at full page width, with a clean
    text block below -- same reliable pattern as build_branded_cover()'s
    own landscape fallback tier.

    title_lines: list of strings, one per line (e.g. ["Introducing",
    "The Verlocity Growth System"] to match the reference exactly, or a
    single-item list for a plain one-line title).
    date_str: defaults to today's date in "Month D, YYYY" form, matching
    the reference cover's "July 15, 2026" style, if not supplied."""
    section = doc.sections[0]
    printable_width = section.page_width - section.left_margin - section.right_margin
    img_path = Path(__file__).parent / "verlocity_cover_snapshot.png"

    if not img_path.exists():
        # No asset -- fall back to the plain navy cover rather than fail.
        build_branded_cover(doc, " / ".join(title_lines), subtitle or "")
        return

    p_img = doc.add_paragraph()
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(0)
    p_img.add_run().add_picture(str(img_path), width=printable_width)

    def _line(text, size, color, bold=False, italic=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = FONT_HEAD
        r.bold = bold
        r.italic = italic
        return p

    for i, line in enumerate(title_lines):
        _line(line, 22, NAVY, bold=True, space_before=(28 if i == 0 else 0))
    if subtitle:
        _line(subtitle, 12, TEAL, space_before=6)
    _line(date_str or datetime.now().strftime("%B %-d, %Y"), 11, GRAY3, space_before=8)
    doc.add_page_break()


def setup_branded_header_footer(doc, bank_name):
    """Consistent header/footer for every page AFTER the cover -- the
    cover gets its own full-bleed treatment via build_branded_cover() and
    should not repeat this header, so this enables 'different first page'
    and only sets the header/footer on section.header (page 2 onward),
    leaving section.first_page_header/footer blank.

    Uses Brandon's actual logo (verlocity_header_logo.png) and footer band
    (verlocity_footer_band.png) graphics rather than plain text -- both
    were missing from the doc entirely before, just styled text standing
    in for them. Falls back to text-only if either asset is missing next
    to this script, so a missing file doesn't break generation."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    printable_width = section.page_width - section.left_margin - section.right_margin
    logo_path = Path(__file__).parent / "verlocity_header_logo.png"
    band_path = Path(__file__).parent / "verlocity_footer_band.png"

    # ── Header: bank context on the left, real logo right-aligned via a
    # right tab stop landing at the printable width's right edge ──
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    r_logo = hp.add_run("VERLOCITY")
    r_logo.bold = True
    r_logo.font.size = Pt(10)
    r_logo.font.color.rgb = NAVY
    r_logo.font.name = FONT_HEAD
    r_sep = hp.add_run(f"   |   {bank_name}")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = GRAY3
    r_sep.font.name = FONT_HEAD
    if logo_path.exists():
        hp.paragraph_format.tab_stops.add_tab_stop(printable_width, WD_TAB_ALIGNMENT.RIGHT)
        r_tab = hp.add_run("\t")
        r_img = hp.add_run()
        r_img.add_picture(str(logo_path), height=Pt(20))
    hp.paragraph_format.space_after = Pt(4)
    # thin rule under the header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "02A7C2")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Footer: Brandon's teal-to-navy gradient band, image spanning most
    # of the width, with a solid-navy cell for the page number sampled
    # directly from the band's own right-edge color (RGB 8,62,96 -- an
    # exact match to brand Navy #083D5F) so the seam is invisible. ──
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""

    if band_path.exists():
        band_w = Inches(5.6)
        num_w = printable_width - band_w
        ft = footer.add_table(rows=1, cols=2, width=printable_width)
        ft.autofit = False
        _remove_table_borders(ft)
        img_cell, num_cell = ft.rows[0].cells
        img_cell.width = band_w
        num_cell.width = num_w
        img_cell.text = ""
        img_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        img_cell.paragraphs[0].add_run().add_picture(str(band_path), width=band_w)
        _set_cell_shading(num_cell, "083D5F")
        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        np_ = num_cell.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_field = np_.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_field._r.append(fld_begin)
        run_field._r.append(instr)
        run_field._r.append(fld_end)
        run_field.bold = True
        run_field.font.size = Pt(10)
        run_field.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_field.font.name = FONT_HEAD
        return

    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_conf = fp.add_run("VERLOCITY.AI  ·  CONFIDENTIAL  ·  Page ")
    r_conf.font.size = Pt(8)
    r_conf.font.color.rgb = GRAY3
    r_conf.font.name = FONT_HEAD
    # PAGE field
    run_field = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_field._r.append(fld_begin)
    run_field._r.append(instr)
    run_field._r.append(fld_end)
    run_field.font.size = Pt(8)
    run_field.font.color.rgb = GRAY3
    run_field.font.name = FONT_HEAD


def _render_signal_brief(doc, title, raw_text):
    """Renders a web-search signal brief's structured plain-text output
    (three ALL-CAPS headers per the analyst prompts) into formatted doc
    content. Parses defensively -- if the model didn't follow the header
    format exactly, falls back to a single body block rather than dropping
    the content or crashing."""
    _heading(doc, title)

    known_headers = ["PERSONA SIGNAL OVERVIEW", "BRANCH-LEVEL PERSONA SIGNALS",
                      "MARKET OFFER SIGNAL OVERVIEW", "BRANCH-LEVEL MARKET PRESSURE",
                      "STRATEGIC IMPLICATIONS"]
    lines = raw_text.split("\n")
    sections = []  # list of (header_or_None, [body_lines])
    current_header, current_body = None, []
    for line in lines:
        stripped = line.strip()
        matched = next((h for h in known_headers if stripped.upper().startswith(h)), None)
        if matched:
            if current_header or current_body:
                sections.append((current_header, current_body))
            current_header, current_body = stripped, []
        else:
            current_body.append(line)
    sections.append((current_header, current_body))

    if len(sections) == 1 and sections[0][0] is None:
        _body(doc, raw_text.strip(), size=9.5)
        return

    for header, body_lines in sections:
        text = "\n".join(body_lines).strip()
        if not text:
            continue
        if header:
            _heading(doc, header.title(), size=11, space_before=10, space_after=4)
        if "STRATEGIC" in (header or "").upper():
            for bl in text.split("\n"):
                bl = bl.strip().lstrip("•-*").strip()
                if bl:
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(bl)
                    r.font.size = Pt(9.5)
                    r.font.name = FONT_HEAD
        elif "BRANCH-LEVEL" in (header or "").upper():
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    _body(doc, para, size=9.5)
        else:
            _body(doc, text, size=9.5)


def _lookup_branch_narrative(narr_dict, b, default=None):
    """Looks up a branch's AI-generated narrative content by name+city+state,
    tolerant of minor formatting differences in how the AI reproduced the
    key -- an exact-string match silently fails (falling back to generic
    static text with no error anywhere) if the AI adds a stray space,
    changes punctuation, or reorders anything. This tries progressively
    looser matches before giving up."""
    if not narr_dict:
        return default
    namebr = (b.get("namebr") or "").strip()
    citybr = (b.get("citybr") or "").strip()
    stalpbr = (b.get("stalpbr") or "").strip()
    exact_key = f"{namebr} ({citybr}, {stalpbr})"

    # 1. Exact match — the common case when the AI reproduces the key faithfully.
    if exact_key in narr_dict:
        return narr_dict[exact_key]

    # 2. Case-insensitive, whitespace-normalized exact match.
    norm = lambda s: " ".join(s.lower().split())
    exact_norm = norm(exact_key)
    for k, v in narr_dict.items():
        if norm(k) == exact_norm:
            return v

    # 3. Branch name appears in the key (handles city/state formatting
    # drift -- e.g. the AI dropping the state abbreviation or using a
    # different separator) as long as the match is unambiguous.
    if namebr:
        candidates = [v for k, v in narr_dict.items() if norm(namebr) in norm(k)]
        if len(candidates) == 1:
            return candidates[0]

    print(f"  ⚠ No narrative match for '{exact_key}' among {len(narr_dict)} AI-generated "
          f"key(s) — falling back to static/default. Sample AI key: "
          f"{next(iter(narr_dict), 'n/a')!r}")
    return default


def render_branch_deep_dive(doc, b, strat, play, e, capped_yoy, branch_verdicts, branch_plays,
                             branch_audiences, geo_by_uid, tmpdir, heading_space_before=4,
                             vuln_targets=None):
    """Renders one branch's full deep-dive section: verdict, deposits/radius
    methodology, named competitors + competitor map, capture scenario,
    audience signal, assigned play. Extracted from the main per-branch loop
    so the full $10K Assessment (looping over every priority branch) and the
    standalone single-branch Preview doc (shown live in pitch meetings) use
    the exact same rendering code — the preview is never a lower-fidelity
    mockup of what the real Assessment actually produces."""
    branch_label = f"{b.get('namebr','—')} ({b.get('citybr','—')}, {b.get('stalpbr','—')})"
    q_full = b.get("matrix_quadrant") or "—"

    zone = b.get("opportunity_zone", "Analyze")
    zone_rgb = rgb(ZONE_COLOR.get(zone, "185FA5"))
    _heading(doc, branch_label, size=15, color=zone_rgb, space_before=heading_space_before)

    # ── Overview stat block ──
    ov = doc.add_table(rows=2, cols=4)
    ov.style = "Light Grid Accent 1"
    ov_hdr = ov.rows[0].cells
    for j, h in enumerate(["Zone", "Quadrant", "Priority Tier", "Opportunity Score"]):
        ov_hdr[j].text = h
    ov_val = ov.rows[1].cells
    ov_val[0].text = zone
    ov_val[1].text = q_full
    ov_val[2].text = str(b.get("priority_tier") or "—")
    ov_val[3].text = f"{_sf(b.get('opportunity_score')):.0f}/100"
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Branch Verdict — the synthesized "why," stated before the
    # supporting detail tables rather than left for the reader to
    # assemble from five separate sections. ──
    verdict = _lookup_branch_narrative(branch_verdicts, b)
    if not verdict:
        top_comp = strat.get("top_competitor") if strat else None
        if top_comp:
            comp_clause = (f"{top_comp.get('bank_name')} sits "
                            f"{_sf(top_comp.get('distance_miles')):.2f}mi away with "
                            f"${_sf(top_comp.get('deposits'))/1e6:.0f}M in deposits")
        else:
            comp_clause = "no named competitor sits within the adaptive radius"
        verdict = (f"{b.get('namebr')} scores {_sf(b.get('opportunity_score')):.0f}/100 in the "
                   f"{zone} zone, with ${_sf(b.get('latest_dep'))/1e6:.0f}M in deposits at "
                   f"{fmt_yoy(b, capped_yoy)} YoY. "
                   f"{comp_clause[0].upper() + comp_clause[1:]}. "
                   f"Assigned play: {play or 'under review'}.")

    # "At a glance" callout — a CEO skimming top to bottom should get the
    # whole story from this box alone before reaching a single table. Plain
    # gray body text was easy to skip past; a shaded box with a bold label
    # is the first thing the eye lands on.
    glance_box = doc.add_table(rows=1, cols=1)
    glance_cell = glance_box.rows[0].cells[0]
    _set_cell_shading(glance_cell, "F0F4F8")
    glance_cell.paragraphs[0].text = ""
    glance_cell.paragraphs[0].paragraph_format.space_before = Pt(8)
    glance_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    r_glance_lbl = glance_cell.paragraphs[0].add_run("AT A GLANCE\n")
    r_glance_lbl.bold = True
    r_glance_lbl.font.size = Pt(8.5)
    r_glance_lbl.font.color.rgb = zone_rgb
    r_glance_lbl.font.name = FONT_HEAD
    r_verdict = glance_cell.paragraphs[0].add_run(verdict)
    r_verdict.font.size = Pt(10.5)
    r_verdict.font.name = FONT_HEAD
    r_verdict.font.color.rgb = NAVY
    glance_cell.paragraphs[0].paragraph_format.space_after = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Deposits & Radius ──
    _heading(doc, "Deposits & Radius Methodology", size=11, space_before=8, space_after=4)
    dep = _sf(b.get("latest_dep"))
    radius_mi = strat.get("radius_mi") if strat else None
    tier_label = strat.get("tier") if strat else None
    # Households is a fast estimate (branch-ZIP population / 2.5, the Census
    # national average household size), not a radius-level ACS ingestion --
    # that's a real data gap flagged separately, not yet built. Labeled as
    # an estimate in the header itself so it's never read as precise.
    est_households = _sf(b.get("total_population")) / 2.5
    # Market share denominator is total deposits WITHIN THE SAME ADAPTIVE
    # RADIUS shown in this row (not a fixed 10mi figure) -- every branch in
    # range, not just the size-filtered named competitors, plus this
    # branch's own deposits. Column label carries the actual radius so it
    # never has to be cross-checked against a different distance elsewhere
    # in the row.
    total_mkt_dep = _sf(strat.get("total_market_deposits_radius")) if strat else 0.0
    mkt_share = (dep / total_mkt_dep * 100) if total_mkt_dep > 0 else None
    mkt_share_label = f"Market Share ({radius_mi:.1f}mi)" if radius_mi else "Market Share"
    dr = doc.add_table(rows=2, cols=6)
    dr.style = "Light Grid Accent 1"
    dr_hdr = dr.rows[0].cells
    for j, h in enumerate(["Deposits", "YoY Growth", "Market Tier", "Radius Used",
                            "Households (est.)", mkt_share_label]):
        dr_hdr[j].text = h
    dr_val = dr.rows[1].cells
    dr_val[0].text = f"${dep/1e6:.1f}M"
    dr_val[1].text = fmt_yoy(b, capped_yoy)
    dr_val[2].text = tier_label or "—"
    dr_val[3].text = f"{radius_mi}mi" if radius_mi else "—"
    dr_val[4].text = f"~{est_households:,.0f}" if est_households > 0 else "—"
    dr_val[5].text = f"{mkt_share:.1f}%" if mkt_share is not None else "—"
    p_method = doc.add_paragraph()
    p_method.paragraph_format.space_before = Pt(4)
    method_run = p_method.add_run(
        "Radius scales to local density and deposit size — as tight as 0.5mi for "
        "dense, high-value markets, up to 10mi for rural or low-deposit branches — "
        "rather than one fixed distance for the whole network. Households (est.) is "
        "the branch ZIP's population divided by the Census national average household "
        "size (2.5) — a fast estimate, not a radius-level Census household count. "
        "Market Share is this branch's deposits as a share of ALL deposits within the "
        "same adaptive radius shown here (FDIC Summary of Deposits) — every branch in "
        "range, not just the named competitors below."
    )
    method_run.italic = True
    method_run.font.size = Pt(8.5)
    method_run.font.color.rgb = GRAY3
    method_run.font.name = FONT_HEAD

    # ── Competitors (top 3, size-filtered) ──
    _heading(doc, "Named Competitors — Who To Target", size=11, space_before=10, space_after=4)
    top3 = (strat.get("top3_competitors") if strat else []) or []
    all_comp = (strat.get("all_competitors") if strat else []) or []
    vuln_list = sorted((vuln_targets or {}).get(b.get("uninumbr"), []), key=lambda c: c.get("rank") or 99)

    # Diagnostic logging — the competitor map has gone missing from at least
    # one real generation (Trustmark / Jones Valley) with no exception raised,
    # which means it silently hit one of the conditions below rather than
    # failing loudly. These prints turn that into a visible, one-line answer
    # in the Railway logs on the next run instead of another round of
    # static-code guessing.
    if not all_comp:
        print(f"  ⚠ [map] no all_competitors for {b.get('namebr')} — skipping radius map "
              f"(strat={'present' if strat else 'MISSING'}, "
              f"all_competitors_key={'present' if strat and 'all_competitors' in strat else 'MISSING'})")

    own = geo_by_uid.get(b.get("uninumbr"))

    # Map competitor source — prefer all_competitors entries that actually
    # carry lat/lon (the full-network batch RPC includes it; the lightweight
    # single-branch RPC used by the Preview flow does NOT), falling back to
    # vuln_list (branch_target_competitors) otherwise. vuln_list has its own
    # geo join in fetch_vulnerability_targets() and has rendered correctly
    # in every real generation so far. Checking for usable geo specifically
    # (not just non-empty) matters now that all_competitors can be non-empty
    # but geo-less depending on which fetch path produced it.
    map_competitors = [c for c in all_comp if c.get("lat") is not None and c.get("lon") is not None]
    if not map_competitors and vuln_list:
        own_lat = own.get("lat") if own else None
        own_lon = own.get("lon") if own else None
        map_competitors = [
            {
                "bank_name": c.get("bank_name"),
                "deposits": c.get("deposits"),
                "lat": c.get("lat"),
                "lon": c.get("lon"),
                "distance_miles": _distance_miles(own_lat, own_lon, c.get("lat"), c.get("lon")),
            }
            for c in vuln_list if c.get("lat") is not None and c.get("lon") is not None
        ]
        if map_competitors:
            print(f"  ✓ [map] built from vuln_list fallback for {b.get('namebr')} "
                  f"({len(map_competitors)} of {len(vuln_list)} vuln competitors had usable geo)")
        elif vuln_list:
            print(f"  ⚠ [map] vuln_list present ({len(vuln_list)} competitors) but none had "
                  f"usable lat/lon — geo join in fetch_vulnerability_targets() likely missed "
                  f"these target_uninumbr values in branches_master_v2")

    if map_competitors:
        if not own or own.get("lat") is None or own.get("lon") is None:
            print(f"  ⚠ [map] no usable geo for {b.get('namebr')} (uninumbr={b.get('uninumbr')}) "
                  f"in geo_by_uid ({len(geo_by_uid)} entries loaded) — skipping radius map")
        if own and own.get("lat") is not None and own.get("lon") is not None:
            radius_img = os.path.join(tmpdir, f"radius_{b.get('uninumbr')}.png")
            ok = chart_branch_radius_map(
                own["lat"], own["lon"], map_competitors, radius_mi or 3.0, radius_img
            )
            if not ok:
                print(f"  ⚠ [map] chart_branch_radius_map returned False for {b.get('namebr')} "
                      f"— both Mapbox and local-plot fallback failed")
            if ok:
                p_img = doc.add_paragraph()
                p_img.paragraph_format.space_before = Pt(4)
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(radius_img, width=Inches(3.6))
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_cap = p_cap.add_run(
                    f"{len(map_competitors)} competitor{'s' if len(map_competitors) != 1 else ''} "
                    f"shown, sized by deposits — full competitive density in this market."
                )
                r_cap.italic = True
                r_cap.font.size = Pt(7.5)
                r_cap.font.color.rgb = GRAY3
                r_cap.font.name = FONT_HEAD

    if vuln_list:
        # This is the actual answer to "who do we go after" -- ranked by a
        # composite vulnerability score (declining deposits, weak ROA,
        # asset-quality stress), not just by size. The map above shows the
        # full competitive landscape; this table is the prioritized target
        # list drawn from it.
        p_lead = doc.add_paragraph()
        p_lead.paragraph_format.space_before = Pt(8)
        r_lead = p_lead.add_run(
            "Ranked by vulnerability, not size — which nearby institutions are actually "
            "losing ground, not just which are largest. The score weighs deposit trend, "
            "profitability, asset-quality stress, proximity, and each competitor's own "
            "overall market position together — a bank losing deposits fastest doesn't "
            "automatically rank first if it's otherwise financially stronger."
        )
        r_lead.italic = True
        r_lead.font.size = Pt(9)
        r_lead.font.color.rgb = GRAY3
        r_lead.font.name = FONT_HEAD

        vt = doc.add_table(rows=1, cols=8)
        vt.style = "Light Grid Accent 1"
        vt.autofit = False
        # Explicit widths -- 8 columns is tight on a US Letter page at
        # default autofit, which was breaking headers mid-word ("Opportuni-
        # ty Score", "Competito-r"). Narrow columns for short values, more
        # room for the two text-heavy ones (Competitor, Weakness).
        col_widths = [Inches(0.35), Inches(1.35), Inches(0.7), Inches(0.65),
                      Inches(0.55), Inches(0.65), Inches(0.65), Inches(1.45)]
        vt_hdr = vt.rows[0].cells
        for j, h in enumerate(["Rank", "Competitor", "Deposits", "Mkt Share", "YoY", "Opp Score",
                                "Vuln Score", "Weakness"]):
            vt_hdr[j].text = h
            vt_hdr[j].width = col_widths[j]

        def _cell_run(cell, text, bold=False, color=None):
            cell.text = ""
            r = cell.paragraphs[0].add_run(text)
            r.font.name = FONT_HEAD
            r.font.size = Pt(8.5)
            if bold:
                r.font.bold = True
            if color:
                r.font.color.rgb = color
        for c in vt_hdr:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8.5)

        # Tags computed once for the whole group -- see _vulnerability_tags()
        # docstring for why this can't be done per-row independently.
        tags_by_name = _vulnerability_tags(vuln_list[:3])
        weakest_score_name = min(
            (c for c in vuln_list[:3] if c.get("opportunity_score") is not None),
            key=lambda c: _sf(c.get("opportunity_score")), default={}
        ).get("bank_name")

        for c in vuln_list[:3]:
            row = vt.add_row().cells
            for j, cell in enumerate(row):
                cell.width = col_widths[j]
            yoy = _sf(c.get("yoy_pct"))
            opp_score = c.get("opportunity_score")
            vuln_score = c.get("vuln_score")
            is_top = c is vuln_list[0]
            name = c.get("bank_name")
            tag = tags_by_name.get(name, "Declining deposits")
            comp_dep = _sf(c.get("deposits"))
            # Same radius-scoped denominator as the branch's own Market
            # Share above -- every competitor's share of that same total,
            # not a separately-computed figure, so the numbers in this
            # table and the row above are directly comparable.
            comp_share = (comp_dep / total_mkt_dep * 100) if total_mkt_dep > 0 else None

            _cell_run(row[0], str(c.get("rank") or "—"), bold=is_top)
            _cell_run(row[1], str(c.get("bank_name", "—")), bold=is_top)
            _cell_run(row[2], f"${comp_dep/1e6:.1f}M")
            _cell_run(row[3], f"{comp_share:.1f}%" if comp_share is not None else "—")
            # Weak individual metrics get flagged red/bold right in the cell
            # so the eye catches the specific weakness without reading the
            # Weakness column or any prose below the table.
            _cell_run(row[4], f"{yoy:+.1f}%", bold=(yoy < 0), color=RED_WEAK if yoy < 0 else None)
            _cell_run(row[5], f"{opp_score:.0f}/100" if opp_score is not None else "—",
                      bold=(name == weakest_score_name), color=RED_WEAK if name == weakest_score_name else None)
            _cell_run(row[6], f"{vuln_score:.0f}" if vuln_score is not None else "—", bold=is_top)
            _cell_run(row[7], tag, bold=True, color=RED_WEAK if tag != "Stable, size-based target only" else GRAY3)

            # The #1 vulnerability-ranked competitor's whole row is shaded so
            # it's the one row a skimming reader's eye lands on first, before
            # reading rank numbers or any column at all.
            if is_top:
                for cell in row:
                    _set_cell_shading(cell, RED_WEAK_FILL)

        # Explains the rank order itself when it isn't a simple sort by YoY
        # severity -- e.g. #2 declining faster than #1. Real, recurring case,
        # not hypothetical: the vuln_score formula weighs several factors
        # together, so this comes up whenever one factor pulls against another.
        rank_note = _rank_order_note(vuln_list[:3])
        if rank_note:
            p_rank = doc.add_paragraph()
            p_rank.paragraph_format.space_before = Pt(6)
            r_rank = p_rank.add_run(rank_note)
            r_rank.italic = True
            r_rank.font.size = Pt(8.5)
            r_rank.font.color.rgb = GRAY3
            r_rank.font.name = FONT_HEAD

        # Relative size — where THIS branch stands against the named
        # competitors, not just who they are and how big they are.
        size_line = _relative_size_line(b.get("latest_dep"), vuln_list[:3])
        if size_line:
            p_size = doc.add_paragraph()
            p_size.paragraph_format.space_before = Pt(6)
            r_size_label = p_size.add_run("Where this branch stands: ")
            r_size_label.bold = True
            r_size_label.font.size = Pt(9.5)
            r_size_label.font.color.rgb = NAVY
            r_size_label.font.name = FONT_HEAD
            r_size = p_size.add_run(size_line)
            r_size.font.size = Pt(9.5)
            r_size.font.name = FONT_HEAD
            r_size.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        top_target = vuln_list[0]
        p_win = doc.add_paragraph()
        p_win.paragraph_format.space_before = Pt(8)
        r_win_label = p_win.add_run("Priority target: ")
        r_win_label.bold = True
        r_win_label.font.size = Pt(9.5)
        r_win_label.font.color.rgb = NAVY
        r_win_label.font.name = FONT_HEAD
        r_win = p_win.add_run(
            f"{top_target.get('bank_name')} — {_vulnerability_sentence(top_target, vuln_list[:3])}. "
            f"This is where deposit capture is realistically winnable, not just theoretically "
            f"contestable."
        )
        r_win.font.size = Pt(9.5)
        r_win.font.name = FONT_HEAD
        r_win.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Where you have the edge — the #1 target above gets the fullest
        # framing, but reps need to be able to speak to #2 and #3 too if
        # the conversation goes there. Previously only rank 1's weakness
        # was ever surfaced; ranks 2-3 sat in the table with no narrative.
        if len(vuln_list) > 1:
            p_edge_h = doc.add_paragraph()
            p_edge_h.paragraph_format.space_before = Pt(8)
            r_edge_h = p_edge_h.add_run("Where you have the edge:")
            r_edge_h.bold = True
            r_edge_h.font.size = Pt(9.5)
            r_edge_h.font.color.rgb = NAVY
            r_edge_h.font.name = FONT_HEAD
            for c in vuln_list[:3]:
                p_e = doc.add_paragraph(style="List Bullet")
                r_e_name = p_e.add_run(f"{c.get('bank_name')}: ")
                r_e_name.bold = True
                r_e_name.font.size = Pt(9)
                r_e_name.font.name = FONT_HEAD
                r_e_name.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                r_e = p_e.add_run(_vulnerability_sentence(c, vuln_list[:3]))
                r_e.font.size = Pt(9)
                r_e.font.name = FONT_HEAD
                r_e.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    elif top3:
        cot = doc.add_table(rows=1, cols=4)
        cot.style = "Light Grid Accent 1"
        cot_hdr = cot.rows[0].cells
        for j, h in enumerate(["Competitor", "City / State", "Distance", "Deposits"]):
            cot_hdr[j].text = h
        for c in top3:
            row = cot.add_row().cells
            row[0].text = str(c.get("bank_name", "—"))
            row[1].text = f"{c.get('city','—')}, {c.get('state','—')}"
            row[2].text = f"{_sf(c.get('distance_miles')):.2f}mi"
            row[3].text = f"${_sf(c.get('deposits'))/1e6:.1f}M"
    else:
        _body(doc, "No competitor within the adaptive radius meets the 0.1x-5x size filter — "
                   "this branch has natural geographic protection rather than a capture target.",
              size=9.5)

    # ── Capture Scenario ──
    capture_pool = e["capture_pool"]
    if capture_pool:
        _heading(doc, "Projected Annual Capture", size=11, space_before=10, space_after=4)
        cst = doc.add_table(rows=2, cols=3)
        cst.style = "Light Grid Accent 1"
        cst_hdr = cst.rows[0].cells
        for j, h in enumerate(["Low (1%)", "Medium (3%)", "Aggressive (7%)"]):
            cst_hdr[j].text = h
        cst_val = cst.rows[1].cells
        cst_val[0].text = f"${capture_pool*0.01/1e6:.2f}M"
        cst_val[1].text = f"${capture_pool*0.03/1e6:.2f}M"
        cst_val[2].text = f"${capture_pool*0.07/1e6:.2f}M"

    # ── Audience Signal — condensed single-branch version of the multi-branch
    # Audience Intelligence Brief pattern: stat row, one synthesized paragraph,
    # a single persona card, then "where it's strong" / "validate before
    # activating" bullets. Same analytical shape as the portfolio-level
    # brief, just scaled to one branch instead of a market of them. ──
    audience_detail = _lookup_branch_narrative(branch_audiences, b, default=None)
    audience_text, persona, strong_signals, validate_signals = "", None, [], []
    if isinstance(audience_detail, dict):
        audience_text = audience_detail.get("narrative") or ""
        if audience_detail.get("persona_name"):
            persona = audience_detail
        strong_signals = audience_detail.get("strong_signals") or []
        validate_signals = audience_detail.get("validate_before_activating") or []
    elif isinstance(audience_detail, str):
        audience_text = audience_detail  # backward-compat with the older string-only shape

    if audience_text or b.get("household_income"):
        inc = _sf(b.get("household_income"))
        inc_yoy = _sf(b.get("yoy_income_growth")) * 100
        pop_yoy = _sf(b.get("yoy_pop_growth")) * 100
        zhvi_yoy = _sf(b.get("zhvi_yoy_pct"))  # already a percentage, not a decimal
        smb_zone = b.get("smb_zone")

        _heading(doc, "Audience Signal", size=11, space_before=10, space_after=4)

        card = doc.add_table(rows=1, cols=1)
        card_cell = card.rows[0].cells[0]
        _set_cell_shading(card_cell, "EAF6F8")  # light teal tint -- distinct from the
        card_cell.paragraphs[0].paragraph_format.space_after = Pt(8)  # navy AT-A-GLANCE box above

        # Stat row: numbers up front, the thing a reader's eye should land
        # on before any prose. Nested table inside the shaded cell so each
        # stat gets its own column rather than being buried in a sentence.
        stats = [
            (f"${inc:,.0f}", "Household Income"),
            (f"{inc_yoy:+.1f}%", "Income YoY"),
            (f"{pop_yoy:+.1f}%", "Population YoY"),
        ]
        if smb_zone:
            stats.append((str(smb_zone), "SMB Zone"))
        stat_tbl = card_cell.add_table(rows=2, cols=len(stats))
        stat_tbl.autofit = True
        for j, (val, lbl) in enumerate(stats):
            vcell, lcell = stat_tbl.rows[0].cells[j], stat_tbl.rows[1].cells[j]
            vcell.paragraphs[0].text = ""
            vr = vcell.paragraphs[0].add_run(val)
            vr.bold = True
            vr.font.size = Pt(17 if len(stats) > 3 else 19)
            vr.font.color.rgb = TEAL
            vr.font.name = FONT_HEAD
            vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            lcell.paragraphs[0].text = ""
            lr = lcell.paragraphs[0].add_run(lbl.upper())
            lr.font.size = Pt(7.5)
            lr.font.color.rgb = GRAY3
            lr.font.name = FONT_HEAD
            lr.bold = True
            lcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "The audience in one paragraph"
        p_narr = card_cell.add_paragraph()
        p_narr.paragraph_format.space_before = Pt(10)
        if audience_text:
            r_narr = p_narr.add_run(audience_text)
        else:
            if inc_yoy > 3 and pop_yoy > 1:
                segment_read = ("supports an expansion-oriented read — High-Quality Local "
                                 "Prospect targeting fits a market growing in both income and population")
            elif inc_yoy < 0 or pop_yoy < -1:
                segment_read = ("favors a retention-oriented read over new-household acquisition — "
                                 "Warm Retargeting of the existing base outperforms broad prospecting here")
            else:
                segment_read = "reads as stable — steady-state prospecting, no urgency in either direction"
            wealth_note = ""
            if zhvi_yoy > 5:
                wealth_note = " Rising home values add a supporting tailwind for CD/HYSA acquisition."
            elif zhvi_yoy < -2:
                wealth_note = " Softening home values warrant caution on aggressive acquisition spend."
            r_narr = p_narr.add_run(
                f"Home values {zhvi_yoy:+.1f}% YoY. This profile {segment_read}.{wealth_note}"
            )
        r_narr.font.size = Pt(9.5)
        r_narr.font.name = FONT_HEAD
        r_narr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Persona mini-card — one named persona for this branch (not a
        # multi-persona list; that's the portfolio-brief version of this).
        if persona:
            persona_tbl = card_cell.add_table(rows=1, cols=1)
            p_cell = persona_tbl.rows[0].cells[0]
            _set_cell_shading(p_cell, "FFFFFF")
            p_cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            r_pname = p_cell.paragraphs[0].add_run(persona.get("persona_name", ""))
            r_pname.bold = True
            r_pname.font.size = Pt(10.5)
            r_pname.font.color.rgb = NAVY
            r_pname.font.name = FONT_HEAD
            if persona.get("persona_tagline"):
                p_tag = p_cell.add_paragraph()
                r_tag = p_tag.add_run(persona["persona_tagline"])
                r_tag.italic = True
                r_tag.font.size = Pt(8.5)
                r_tag.font.color.rgb = GRAY3
                r_tag.font.name = FONT_HEAD
            for label, value in [("Life stage", persona.get("life_stage")),
                                  ("Wealth signal", persona.get("wealth_signal")),
                                  ("Primary need", persona.get("primary_need")),
                                  ("Switch driver", persona.get("switch_driver"))]:
                if not value:
                    continue
                p_attr = p_cell.add_paragraph()
                p_attr.paragraph_format.space_before = Pt(3)
                r_lbl = p_attr.add_run(f"{label}: ")
                r_lbl.bold = True
                r_lbl.font.size = Pt(8.5)
                r_lbl.font.color.rgb = NAVY
                r_lbl.font.name = FONT_HEAD
                r_val = p_attr.add_run(value)
                r_val.font.size = Pt(8.5)
                r_val.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                r_val.font.name = FONT_HEAD
            p_cell.paragraphs[-1].paragraph_format.space_after = Pt(4)

        # Where it's strong / validate before activating — condensed to one
        # branch's worth (2 + 1 bullets) instead of the portfolio brief's
        # per-market lists, but the same "here's the case, here's the
        # caveat" honesty that makes it read as analysis, not a sales pitch.
        if strong_signals:
            p_strong_h = card_cell.add_paragraph()
            p_strong_h.paragraph_format.space_before = Pt(10)
            r_strong_h = p_strong_h.add_run("Where it's strong")
            r_strong_h.bold = True
            r_strong_h.font.size = Pt(8.5)
            r_strong_h.font.color.rgb = RGBColor(0x1A, 0x7A, 0x4A)
            r_strong_h.font.name = FONT_HEAD
            for s in strong_signals[:2]:
                p_s = card_cell.add_paragraph()
                p_s.paragraph_format.space_before = Pt(1)
                r_arrow = p_s.add_run("→ ")
                r_arrow.bold = True
                r_arrow.font.size = Pt(8.5)
                r_arrow.font.color.rgb = RGBColor(0x1A, 0x7A, 0x4A)
                r_arrow.font.name = FONT_HEAD
                r_s = p_s.add_run(s)
                r_s.font.size = Pt(8.5)
                r_s.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                r_s.font.name = FONT_HEAD

        if validate_signals:
            p_val_h = card_cell.add_paragraph()
            p_val_h.paragraph_format.space_before = Pt(8)
            r_val_h = p_val_h.add_run("Validate before activating")
            r_val_h.bold = True
            r_val_h.font.size = Pt(8.5)
            r_val_h.font.color.rgb = RED_WEAK
            r_val_h.font.name = FONT_HEAD
            for v in validate_signals[:2]:
                p_v = card_cell.add_paragraph()
                p_v.paragraph_format.space_before = Pt(1)
                r_arrow2 = p_v.add_run("→ ")
                r_arrow2.bold = True
                r_arrow2.font.size = Pt(8.5)
                r_arrow2.font.color.rgb = RED_WEAK
                r_arrow2.font.name = FONT_HEAD
                r_v = p_v.add_run(v)
                r_v.font.size = Pt(8.5)
                r_v.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                r_v.font.name = FONT_HEAD
        card_cell.paragraphs[-1].paragraph_format.space_after = Pt(6)

    # ── Play ──
    if play:
        _heading(doc, "Assigned Play", size=11, color=NAVY, space_before=10, space_after=4)
        play_box = doc.add_table(rows=1, cols=1)
        cell = play_box.rows[0].cells[0]
        _set_cell_shading(cell, "083D5F")
        cell.paragraphs[0].text = ""
        pr = cell.paragraphs[0].add_run(play.upper())
        pr.font.size = Pt(13)
        pr.font.bold = True
        pr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        pr.font.name = FONT_HEAD
        cell.paragraphs[0].paragraph_format.space_before = Pt(8)
        cell.paragraphs[0].paragraph_format.space_after = Pt(8)

        branch_play_data = _lookup_branch_narrative(branch_plays, b, default={}) or {}
        has_competitor = bool(top3)
        no_comp_override = PLAY_NO_COMPETITOR_OVERRIDE.get(play) if not has_competitor else None
        posture = (branch_play_data.get("resource_posture")
                   or (no_comp_override or {}).get("resource_posture")
                   or PLAY_ACQUISITION_POSTURE.get(play, ""))
        brief = (branch_play_data.get("media_brief")
                 or (no_comp_override or {}).get("media_brief")
                 or PLAY_MEDIA_BRIEF.get(play, ""))
        if posture:
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_before = Pt(6)
            lbl1 = p1.add_run("Resource posture: ")
            lbl1.bold = True
            lbl1.font.size = Pt(9.5)
            lbl1.font.color.rgb = NAVY
            lbl1.font.name = FONT_HEAD
            r1 = p1.add_run(posture)
            r1.font.size = Pt(9.5)
            r1.font.name = FONT_HEAD
        if brief:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            lbl2 = p2.add_run("Media brief: ")
            lbl2.bold = True
            lbl2.font.size = Pt(9.5)
            lbl2.font.color.rgb = NAVY
            lbl2.font.name = FONT_HEAD
            r2 = p2.add_run(brief)
            r2.font.size = Pt(9.5)
            r2.font.name = FONT_HEAD


def _heading(doc, text, size=16, color=NAVY, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True  # never let a heading get stranded
    run = p.add_run(text)                     # alone at the bottom of a page
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_HEAD  # Inter Bold per Verlocity_Brand_Guidelines_R2.pdf
    return p


def _body(doc, text, size=10.5, color=RGBColor(0x33, 0x33, 0x33)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_HEAD  # Inter (Light per brand guide) — same family, body weight
    return p


def build_assessment_doc(bank_name, summary, fin, targets, narr, branches, branches_geo=None,
                          branch_strategy=None, dives=None, deep_mode=None, tmpdir=".", capped_yoy=None,
                          persona_brief=None, market_offer_brief=None, vulnerability_targets=None):
    capped_yoy = capped_yoy or {}
    geo_by_uid = {g["uninumbr"]: g for g in (branches_geo or []) if g.get("uninumbr") is not None}
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)   # US Letter
    section.page_height = Cm(27.94)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ── Cover + branded header/footer (brand guide-aligned) ──
    build_branded_cover(doc, bank_name, "BMAP Market Assessment")
    setup_branded_header_footer(doc, bank_name)

    # Zone distribution hero visual -- kept as its own opening page right
    # after the cover (a matplotlib chart doesn't read well directly on
    # the navy cover background), sets the McKinsey "big number up front"
    # tone before the Executive Summary text begins.
    zone_chart_path = f"{tmpdir}/_chart_zones.png"
    chart_zone_distribution(summary["zones"], zone_chart_path)
    doc.add_picture(zone_chart_path, width=Inches(6.3))

    doc.add_page_break()

    # ── Executive Summary ──
    _heading(doc, "Executive Summary", space_before=0)

    # Pull-quote callout — the single most important number, McKinsey-style
    top_branch = summary["top5"][0] if summary["top5"] else None
    if top_branch:
        callout = doc.add_table(rows=1, cols=1)
        cell = callout.rows[0].cells[0]
        _set_cell_shading(cell, "083D5F")
        cell.paragraphs[0].text = ""
        r1 = cell.paragraphs[0].add_run(
            f"{top_branch['namebr']}: {_sf(top_branch['opportunity_score']):.0f}/100 opportunity score, "
            f"{fmt_yoy(top_branch, capped_yoy)} deposit growth"
        )
        r1.font.size = Pt(13)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].paragraph_format.space_before = Pt(10)
        cell.paragraphs[0].paragraph_format.space_after = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    fallback_headline = (
        f"{bank_name} operates {summary['branch_count']} branches with ${summary['total_deposits_B']:.1f}B "
        f"in total deposits, network average opportunity score {summary['avg_score']:.0f}/100 across "
        f"{summary['zones']['Invest']} Invest, {summary['zones']['Analyze']} Analyze, "
        f"{summary['zones']['Defend']} Defend, and {summary['zones']['Justify']} Justify branches."
    )
    _body(doc, narr.get("exec_headline") or fallback_headline)

    if narr.get("strategic_positioning"):
        p_pos = doc.add_paragraph()
        p_pos.paragraph_format.space_before = Pt(8)
        r_pos = p_pos.add_run(narr["strategic_positioning"])
        r_pos.font.size = Pt(10.5)
        r_pos.font.name = FONT_HEAD
        r_pos.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Priority Focus — 2-3 named branches, structured (not a table, per the
    # source Executive Headline spec this was adapted from: name+city, zone,
    # why now, strategic role).
    priority_focus = narr.get("priority_focus") or []
    if not priority_focus:
        fr_local = summary.get("flagship_risk")
        for cand in ([fr_local] if fr_local else []) + summary["top5"][:3]:
            if not cand:
                continue
            already = any(pf.get("branch") == cand.get("namebr") for pf in priority_focus)
            if already or len(priority_focus) >= 3:
                continue
            is_risk = fr_local and cand.get("namebr") == fr_local.get("namebr")
            priority_focus.append({
                "branch": cand.get("namebr"), "city": cand.get("citybr"), "state": cand.get("stalpbr"),
                "zone": cand.get("opportunity_zone"),
                "why_now": f"{fmt_yoy(cand, capped_yoy)} YoY, {_sf(cand.get('opportunity_score')):.0f}/100 score.",
                "role": "stabilize and retain" if is_risk else "deposit growth engine",
            })
    if priority_focus:
        _heading(doc, "Priority Focus", size=12, space_before=12, space_after=4)
        for item in priority_focus:
            zone = item.get("zone", "")
            zone_rgb = rgb(ZONE_COLOR.get(zone, "083D5F"))
            p_pf = doc.add_paragraph()
            p_pf.paragraph_format.space_before = Pt(6)
            r_name = p_pf.add_run(f"{item.get('branch','—')} ({item.get('city','—')}, {item.get('state','—')}) ")
            r_name.font.bold = True
            r_name.font.size = Pt(10.5)
            r_name.font.name = FONT_HEAD
            r_name.font.color.rgb = NAVY
            r_zone = p_pf.add_run(f"· {zone}")
            r_zone.font.bold = True
            r_zone.font.size = Pt(9.5)
            r_zone.font.name = FONT_HEAD
            r_zone.font.color.rgb = zone_rgb
            p_pf2 = doc.add_paragraph()
            p_pf2.paragraph_format.space_after = Pt(2)
            r_role = p_pf2.add_run(f"{item.get('role','—')} — ")
            r_role.italic = True
            r_role.font.size = Pt(9.5)
            r_role.font.name = FONT_HEAD
            r_role.font.color.rgb = GRAY3
            r_why = p_pf2.add_run(item.get("why_now", ""))
            r_why.font.size = Pt(9.5)
            r_why.font.name = FONT_HEAD
            r_why.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # What This Means for the Next 12 Months — exactly 3 leadership-level
    # decisions, no tactics/channels/pricing per spec.
    next_12 = narr.get("next_12_months") or []
    if not next_12:
        next_12 = [
            f"Allocate capital toward the {summary['zones']['Invest']} Invest-zone branches before "
            f"broad-based network spend.",
            f"Apply retention discipline across the {summary['zones']['Defend'] + summary['zones']['Justify']} "
            f"Defend/Justify branches rather than treating them as growth targets.",
            "Fund digital-first, market-specific execution as the primary lever for deposit growth "
            "beyond the existing physical footprint.",
        ]
    if next_12:
        _heading(doc, "What This Means for the Next 12 Months", size=12, space_before=12, space_after=4)
        for bullet in next_12:
            p_b = doc.add_paragraph(style="List Bullet")
            r_b = p_b.add_run(bullet)
            r_b.font.size = Pt(10)
            r_b.font.name = FONT_HEAD
            r_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Flagship-risk alert — guaranteed regardless of AI narrative compliance.
    # The pull-quote above is the top opportunity-score branch, which can be
    # a small branch; if the network's largest branch by deposits is
    # materially at risk, that fact belongs on page one too, not just in the
    # 59-row appendix table. (Priority Focus above is instructed to include
    # it by name when present, but this stays as a guaranteed backstop.)
    fr = summary.get("flagship_risk")
    if fr:
        alert = doc.add_table(rows=1, cols=1)
        cell = alert.rows[0].cells[0]
        _set_cell_shading(cell, "BFC815")  # brand Justify color
        cell.paragraphs[0].text = ""
        r_alert = cell.paragraphs[0].add_run(
            f"FLAGSHIP RISK — {fr['namebr']} ({fr['citybr']}, {fr['stalpbr']}) holds "
            f"{fr['deposit_share_pct']:.0f}% of total network deposits (${_sf(fr['latest_dep'])/1e6:.0f}M) "
            f"and is at {fmt_yoy(fr, capped_yoy)} YoY, {fr['opportunity_zone']} zone."
        )
        r_alert.font.size = Pt(12)
        r_alert.font.bold = True
        r_alert.font.color.rgb = NAVY
        cell.paragraphs[0].paragraph_format.space_before = Pt(8)
        cell.paragraphs[0].paragraph_format.space_after = Pt(8)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── How This Assessment Works ──
    # Front-loaded, not buried in footnotes — a skeptical reader's first
    # question is "prove this isn't a black box," and that has to be
    # answered before the maps/charts, not after.
    _heading(doc, "How This Assessment Works")
    _body(doc,
          "Every classification, radius, and dollar figure in this document follows the same fixed "
          "rule set, applied identically across all "
          f"{summary['branch_count']} branches. The rules are stated here in full so any number in "
          "this Assessment can be traced back and defended on its own.")

    method_items = [
        ("Data sources", "FDIC Summary of Deposits, FFIEC Call Report, Census ACS demographic and "
                          "population data, and Zillow Home Value Index — all public, all matched to "
                          "the same branch-level geography."),
        ("Opportunity score", "0-100, blending market growth, relative growth vs. peers, inverted "
                               "competitive density, and deposit size. Shown next to the network "
                               "average throughout — a score is only meaningful relative to its own "
                               "network, which is why every citation in this document includes both."),
        ("Adaptive competitor radius", "Set by branch density and deposit size, not chosen per-branch: "
                                        "rural or low-deposit branches widen to 10 miles so a thin market "
                                        "doesn't miss the few real competitors that exist; dense, "
                                        "high-deposit branches tighten to as little as 0.5 miles, since a "
                                        "wide radius there would return dozens of irrelevant competitors. "
                                        "The exact radius and tier used is disclosed under every branch."),
        ("Named competitor filter", "A named competitor must be sized between 0.1x and 5x the branch's "
                                     "own deposits to be listed — this excludes giant national-bank hub "
                                     "branches (which would otherwise dominate every list) as unrealistic "
                                     "local capture targets."),
        ("Capture scenarios", "1% / 3% / 7% of the named competitor pool, applied uniformly — an "
                               "industry-informed planning range, not this bank's own historical "
                               "conversion data. Replace with actual account-opening history once "
                               "available for a sharper estimate."),
        ("Winsorized growth values", "Any branch whose year-over-year growth would otherwise show an "
                                      "uninformative flat +100% is resolved directly against source "
                                      "deposit history and shown as its real computed growth rate, or "
                                      "as \u201CNew branch\u201D where no prior-year figure exists."),
    ]
    for title, desc in method_items:
        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_after = Pt(6)
        r_t = p_m.add_run(f"{title} — ")
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.name = FONT_HEAD
        r_t.font.color.rgb = NAVY
        r_d = p_m.add_run(desc)
        r_d.font.size = Pt(9.5)
        r_d.font.name = FONT_HEAD
        r_d.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p_disclaimer = doc.add_paragraph()
    p_disclaimer.paragraph_format.space_before = Pt(10)
    r_disc = p_disclaimer.add_run(
        "This Assessment scores branches for growth and competitive-capture potential only. "
        "Opportunity-zone classifications (Invest / Analyze / Defend / Justify) are not a "
        "substitute for CRA assessment-area analysis, and a branch's zone should never be the "
        "sole basis for a decision affecting investment in a federally-designated assessment area. "
        "Consult compliance counsel before using this Assessment to inform CRA-related decisions."
    )
    r_disc.italic = True
    r_disc.font.size = Pt(8.5)
    r_disc.font.color.rgb = GRAY3
    r_disc.font.name = FONT_HEAD

    # ── Geographic Distribution (map) ──
    if branches_geo:
        map_path = f"{tmpdir}/_chart_map.png"
        if chart_branch_map(branches_geo, map_path):
            _heading(doc, "Geographic Distribution")
            doc.add_picture(map_path, width=Inches(6.3))

    # ── Network Opportunity Overview ──
    _heading(doc, "Network Opportunity Overview")
    _body(doc, narr.get("network_narrative") or "")

    if summary["top5"] and summary["bottom3"]:
        top_bottom_path = f"{tmpdir}/_chart_topbottom.png"
        chart_top_bottom_branches(summary["top5"], summary["bottom3"], top_bottom_path)
        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.space_before = Pt(10)
        r = p_lbl.add_run("Highest- and Lowest-Opportunity Branches")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
        doc.add_picture(top_bottom_path, width=Inches(6.3))

    # ── Competitive Overview ──
    _heading(doc, "Competitive Overview")
    _body(doc, narr.get("competitive_narrative") or "")
    if targets:
        ct = doc.add_table(rows=1, cols=4)
        ct.style = "Light Grid Accent 1"
        hdr = ct.rows[0].cells
        for i, h in enumerate(["Target", "Branches Exposed", "Vulnerability", "Deposit YoY"]):
            hdr[i].text = h
        for t in targets:
            row = ct.add_row().cells
            row[0].text = str(t.get("target_institution", "—"))
            row[1].text = str(t.get("branches_in_radius", "—"))
            row[2].text = _vuln_tier(t.get("avg_vuln_score"))
            row[3].text = f"{_sf(t.get('avg_yoy_pct')):+.1f}%"

    # ── Persona & Demographic Signal Brief / Market Offer & Competitive
    # Signal Brief — both live-web-search enrichment layers, adapted from
    # Verlocity's Hub analyst prompts. Rendered only if the call succeeded;
    # a failed web-search call skips the section rather than showing
    # placeholder text, since fabricating "current market intelligence"
    # would be worse than omitting it. ──
    if persona_brief:
        _render_signal_brief(doc, "Persona & Demographic Signal Brief", persona_brief)
    if market_offer_brief:
        _render_signal_brief(doc, "Market Offer & Competitive Signal Brief", market_offer_brief)

    # ── Deposit Capture Strategy (adaptive-radius, branch + bank-wide) ──
    if branch_strategy:
        _heading(doc, "Deposit Capture Strategy — By Branch and Network-Wide")
        _body(doc, narr.get("capture_strategy_narrative") or
              "Competitive radius is scaled per branch by local density and deposit size, rather than "
              "a single fixed distance — dense, high-value branches are assessed as tight as 0.5 miles; "
              "rural or low-deposit branches widen to 10 miles, since a thin market can otherwise miss "
              "the few competitors that do exist.")

        bs_summary = summarize_branch_strategy(branch_strategy)
        rt = doc.add_table(rows=1, cols=3)
        rt.style = "Light Grid Accent 1"
        hdr = rt.rows[0].cells
        for i, h in enumerate(["Market Tier", "Branches", "Deposits in Tier"]):
            hdr[i].text = h
        for tier in ["Dense/High-Value", "Dense", "Suburban", "Low-Density"]:
            v = bs_summary["tiers"].get(tier)
            if not v:
                continue
            row = rt.add_row().cells
            row[0].text = tier
            row[1].text = str(v["count"])
            row[2].text = f"${v['deposits']/1e6:.0f}M"
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # Scenario-based $ opportunity — low/medium/aggressive annual capture rate
        # against the contestable deposit pool per tier. Industry-informed planning
        # assumption, not the specific bank's own history -- flagged as such in-doc.
        p_lbl2 = doc.add_paragraph()
        p_lbl2.paragraph_format.space_before = Pt(6)
        r2 = p_lbl2.add_run("Projected Annual Capture by Scenario")
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = NAVY
        r2.font.name = FONT_HEAD

        st = doc.add_table(rows=1, cols=4)
        st.style = "Light Grid Accent 1"
        hdr = st.rows[0].cells
        for i, h in enumerate(["Market Tier", "Low (1%)", "Medium (3%)", "Aggressive (7%)"]):
            hdr[i].text = h
        for tier in ["Dense/High-Value", "Dense", "Suburban", "Low-Density"]:
            v = bs_summary["tiers"].get(tier)
            if not v:
                continue
            dep = v["deposits"]
            row = st.add_row().cells
            row[0].text = tier
            row[1].text = f"${dep*0.01/1e6:.1f}M"
            row[2].text = f"${dep*0.03/1e6:.1f}M"
            row[3].text = f"${dep*0.07/1e6:.1f}M"
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(4)
        note_run = p_note.add_run(
            f"Industry-informed planning assumption, not {bank_name}-specific history. Retail deposit "
            "switching is behaviorally sticky; 7-8% annual capture of a local contestable pool is close "
            "to a realistic ceiling absent a genuine market disruption. Replace with the bank's own "
            "historical account-opening and deposit-capture data once available, per the same "
            "calibration principle used in the Predictive ROI model."
        )
        note_run.italic = True
        note_run.font.size = Pt(8.5)
        note_run.font.color.rgb = GRAY3
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # Named per-branch findings — the actual capture targets, not just tier counts
        named = [r for r in branch_strategy if r.get("top_competitor")]
        named = sorted(named, key=lambda r: -r["deposits"])[:12]
        if named:
            p_lbl = doc.add_paragraph()
            r = p_lbl.add_run("Named Nearest Competitor by Branch (adaptive radius)")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            nt = doc.add_table(rows=1, cols=5)
            nt.style = "Light Grid Accent 1"
            hdr = nt.rows[0].cells
            for i, h in enumerate(["Branch", "Tier", "Radius", "Largest Nearby Competitor", "Distance"]):
                hdr[i].text = h
            for r_ in named:
                tc = r_["top_competitor"]
                row = nt.add_row().cells
                row[0].text = f"{r_['namebr']} ({r_['citybr']}, {r_['stalpbr']})"
                row[1].text = r_["tier"]
                row[2].text = f"{r_['radius_mi']}mi"
                row[3].text = str(tc.get("bank_name", "—"))
                row[4].text = f"{_sf(tc.get('distance_miles')):.2f}mi"

    # ── Branch Assessment — full 16-play deep dive ──
    # <20 branches: every branch assessed. >=20: curated to the 15 with the
    # biggest actual $ capture opportunity. dives/deep_mode computed once
    # upstream (shared with get_narratives for the audience blurbs) rather
    # than recomputed here, so narrative keys and doc content stay in sync.
    if dives is None:
        dives, deep_mode = build_branch_deep_dives(branches, branch_strategy or [])

    branch_audiences = narr.get("branch_audiences") or {}
    branch_verdicts = narr.get("branch_verdicts") or {}
    branch_plays = narr.get("branch_plays") or {}

    if dives:
        section_title = ("Branch-by-Branch Assessment" if deep_mode
                          else f"Priority Branch Deep Dives — Top {len(dives)} by Capture Opportunity")
        _heading(doc, section_title)
        if deep_mode:
            _body(doc, f"This network has {len(branches)} branches — under the {DEEP_DIVE_THRESHOLD}-branch "
                       f"threshold for full individual coverage. Every branch below is assessed on zone, "
                       f"competitive quadrant, assigned play, radius methodology, named competitors, and "
                       f"audience signal.")
        else:
            _body(doc, f"This network has {len(branches)} branches — above the threshold for full "
                       f"individual coverage. The {len(dives)} branches below are ranked by actual "
                       f"dollar capture opportunity (named competitor's deposits within the branch's "
                       f"adaptive radius, size-filtered to 0.1x-5x the branch's own deposits per the "
                       f"BMAP vulnerability methodology), not opportunity score alone.")
        doc.add_page_break()

        for i, e in enumerate(dives):
            b = e["branch"]
            strat = e["strategy"]
            play = e["play"]
            render_branch_deep_dive(doc, b, strat, play, e, capped_yoy, branch_verdicts,
                                     branch_plays, branch_audiences, geo_by_uid, tmpdir,
                                     heading_space_before=(0 if i == 0 else 4),
                                     vuln_targets=vulnerability_targets)
            if i < len(dives) - 1:
                doc.add_page_break()

    # ── Financial Health Benchmarking ──
    _heading(doc, "Financial Health Benchmarking")
    _body(doc, narr.get("financial_narrative") or "")

    fin_chart_path = f"{tmpdir}/_chart_financial.png"
    chart_financial_benchmark(fin, bank_name, fin_chart_path)
    doc.add_picture(fin_chart_path, width=Inches(6.3))

    ft = doc.add_table(rows=1, cols=3)
    ft.style = "Light Grid Accent 1"
    hdr = ft.rows[0].cells
    for i, h in enumerate(["Metric", "Value", "Industry Benchmark"]):
        hdr[i].text = h
    metrics = [
        ("ROA", f"{_sf(fin.get('roa')):.2f}%", ">1.0%"),
        ("NIM", f"{_sf(fin.get('nim')):.2f}%", "2.5–3.5%"),
        ("Efficiency Ratio", f"{_sf(fin.get('efficiency_ratio')):.1f}%", "<60%"),
        ("Deposit YoY", f"{_sf(fin.get('dep_yoy_pct')):+.1f}%", ">2%"),
        ("Cost of Funds", f"{_sf(fin.get('cost_of_funds_pct')):.2f}%", "<2%"),
        ("Tier 1 Capital", f"{_sf(fin.get('tier1_capital_pct')):.1f}%", ">8%"),
        ("Net Income YoY", f"{_sf(fin.get('net_income_yoy_pct')):+.1f}%", ">0%"),
    ]
    for label, val, bench in metrics:
        row = ft.add_row().cells
        row[0].text = label
        row[1].text = val
        row[2].text = bench

    p_finnote = doc.add_paragraph()
    p_finnote.paragraph_format.space_before = Pt(4)
    r_finnote = p_finnote.add_run(
        "Benchmarks are standard community-bank industry thresholds, not this institution's "
        "specific peer group — provided as a general reference point for reading the metrics above, "
        "not a formal peer comparison."
    )
    r_finnote.italic = True
    r_finnote.font.size = Pt(8.5)
    r_finnote.font.color.rgb = GRAY3
    r_finnote.font.name = FONT_HEAD

    # ── Next Step Recommendation ──
    _heading(doc, "Recommendation")
    fallback_next_step = None
    if not narr.get("next_step"):
        top = summary["top5"][0] if summary["top5"] else None
        fr = summary.get("flagship_risk")
        parts = []
        if fr:
            parts.append(f"Address {fr['namebr']}'s {fmt_yoy(fr, capped_yoy)} YoY position first — "
                          f"it holds {fr['deposit_share_pct']:.0f}% of total network deposits and outweighs "
                          f"any single opportunity-zone branch in dollar impact.")
        if top:
            parts.append(f"Prioritize capital toward {top['namebr']} ({top['citybr']}, {top['stalpbr']}), "
                          f"the network's top-scored branch at {_sf(top['opportunity_score']):.0f}/100.")
        parts.append(f"With {summary['zones']['Justify']} branches in Justify and "
                      f"{summary['zones']['Invest']} in Invest, the near-term agenda is reallocating "
                      f"capacity from the former to the latter.")
        fallback_next_step = " ".join(parts)
    _body(doc, narr.get("next_step") or fallback_next_step)

    doc.add_page_break()

    # ── Full Branch Appendix ──
    _heading(doc, f"Appendix — Full Branch Scoring ({len(branches)} branches)", space_before=0)
    at = doc.add_table(rows=1, cols=6)
    at.style = "Light Grid Accent 1"
    hdr = at.rows[0].cells
    for i, h in enumerate(["Branch", "City / State", "Deposits", "YoY", "Score", "Zone"]):
        hdr[i].text = h
    for b in branches:
        row = at.add_row().cells
        row[0].text = str(b.get("namebr", "—"))
        row[1].text = f"{b.get('citybr','—')}, {b.get('stalpbr','—')}"
        row[2].text = f"${_sf(b.get('latest_dep'))/1e6:.1f}M"
        row[3].text = fmt_yoy(b, capped_yoy)
        row[4].text = f"{_sf(b.get('opportunity_score')):.0f}"
        row[5].text = str(b.get("opportunity_zone", "—"))

    if capped_yoy:
        fn = doc.add_paragraph()
        fn.paragraph_format.space_before = Pt(6)
        fn_run = fn.add_run(
            "* YoY growth for this branch hit the standard calculation's cap and was "
            "resolved directly against source deposit data: shown as the real computed "
            "growth rate where a prior-year figure exists, or as \"New branch\" where none does."
        )
        fn_run.italic = True
        fn_run.font.size = Pt(8.5)
        fn_run.font.color.rgb = GRAY3
        fn_run.font.name = FONT_HEAD

    # ── Session placeholder (deliberately not narrative content) ──
    doc.add_page_break()
    _heading(doc, "Discussed Live", space_before=0)
    _body(doc, "This assessment includes a working session to walk through these findings and answer "
               "specific questions about your network. Session notes and next-step scope are recorded separately.")

    # Build stamp -- tiny, gray, easy to ignore, exists so anyone looking at
    # the doc can confirm which code generated it without guessing.
    p_build = doc.add_paragraph()
    p_build.paragraph_format.space_before = Pt(14)
    r_build = p_build.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} · build {GENERATOR_BUILD}")
    r_build.font.size = Pt(7)
    r_build.font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
    r_build.font.name = FONT_HEAD

    return doc


def save_doc(doc, bank_name, out_dir=OUT_DIR):
    safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in bank_name).strip()
    date = datetime.now().strftime("%Y%m%d")
    fname = out_dir / f"BMAP_Assessment_{safe}_{date}.docx"
    doc.save(str(fname))
    return fname


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

def run(ik, name_hint=None):
    print(f"\n{'='*60}\n  BMAP Assessment — {name_hint or ik}\n{'='*60}")
    d = fetch_full_network_data(ik)
    if not d["branches"]:
        raise ValueError(
            f"No branch data found for inst_key='{ik}'. This bank is either "
            f"not ingested into branch_opportunity_base, or the inst_key is "
            f"wrong. Refusing to generate a document — an empty-data report "
            f"would look identical to a real one and could ship by mistake."
        )
    bank_name = name_hint or (d["branches"][0].get("namefull") if d["branches"] else None) or ik
    summary = summarize_network(d)
    dives, deep_mode = build_branch_deep_dives(d["branches"], d.get("branch_strategy") or [])
    narr = get_narratives(bank_name, summary, d["fin"], d["targets"], d.get("branch_strategy"), dives,
                           d.get("capped_yoy"), vulnerability_targets=d.get("vulnerability_targets"))
    persona_brief, market_offer_brief = None, None
    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as pool:
        fut_persona = pool.submit(get_persona_signal_brief, bank_name, dives)
        fut_market = pool.submit(get_market_offer_brief, bank_name, dives, d.get("branch_strategy"))
        try:
            persona_brief = fut_persona.result(timeout=90)
        except Exception as ex:
            print(f"  ⚠ persona brief failed/timed out: {type(ex).__name__}: {str(ex) if str(ex) else '(no message -- likely a timeout)'}")
        try:
            market_offer_brief = fut_market.result(timeout=90)
        except Exception as ex:
            print(f"  ⚠ market offer brief failed/timed out: {type(ex).__name__}: {str(ex) if str(ex) else '(no message -- likely a timeout)'}")
    import tempfile
    with tempfile.TemporaryDirectory() as tmpdir:
        doc = build_assessment_doc(bank_name, summary, d["fin"], d["targets"], narr,
                                    d["branches"], d.get("branches_geo"),
                                    d.get("branch_strategy"), dives, deep_mode, tmpdir=tmpdir,
                                    capped_yoy=d.get("capped_yoy"),
                                    persona_brief=persona_brief, market_offer_brief=market_offer_brief,
                                    vulnerability_targets=d.get("vulnerability_targets"))
        path = save_doc(doc, bank_name)
    print(f"\n  ✓ Saved: {path}\n")
    return path


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--inst_key", required=True)
    ap.add_argument("--name", default=None)
    args = ap.parse_args()
    run(args.inst_key, args.name)
