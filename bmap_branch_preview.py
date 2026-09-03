"""
bmap_branch_preview.py — Single-Branch Deep Dive Preview

Purpose: a LIVE PITCH-MEETING tool, distinct from both:
  - bmap_snapshot.py    (automated cold/warm prospecting outreach)
  - bmap_assessment_doc.py (the paid $10K full-network deliverable)

This shows a prospect ONE branch's full deep-dive treatment, at the exact
same depth and fidelity as the paid Assessment, to demonstrate what the
full network engagement produces. It is NOT a simplified mockup: it
literally reuses render_branch_deep_dive() from bmap_assessment_doc.py,
the same function that renders every branch in the real paid deliverable.
If that function's output quality changes (better or worse), this preview
changes with it automatically -- there is no separate, divergent code path
to keep in sync.

Usage:
    python bmap_branch_preview.py --inst_key bank_463735 --name "Hancock Whitney Bank"
    python bmap_branch_preview.py --inst_key bank_463735 --branch "Gulfport Main Branch"

If --branch is omitted, the most compelling branch is auto-selected:
flagship risk first (if one exists), otherwise the top opportunity-score
branch -- so a sales rep never has to know the network well enough to
pick a branch themselves.
"""

import os
import sys
import argparse
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Reuses bmap_assessment_doc.py's actual rendering code, styling constants,
# and data-fetch functions -- co-located in the same repo/deploy directory.
import bmap_assessment_doc as bad


def pick_preview_branch(summary, branches, branch_name=None):
    """Selects which branch to preview. Explicit --branch always wins (exact
    or case-insensitive substring match on namebr). Otherwise auto-picks the
    most compelling story available: the flagship-risk branch if one exists
    (a bank's biggest branch in real trouble is the most attention-grabbing
    opener in a pitch), falling back to the top opportunity-score branch."""
    if branch_name:
        needle = branch_name.strip().lower()
        exact = [b for b in branches if (b.get("namebr") or "").strip().lower() == needle]
        if exact:
            return exact[0]
        partial = [b for b in branches if needle in (b.get("namebr") or "").lower()]
        if len(partial) == 1:
            return partial[0]
        if len(partial) > 1:
            names = ", ".join(b.get("namebr") for b in partial)
            raise ValueError(f"'{branch_name}' matches multiple branches: {names}. Be more specific.")
        raise ValueError(f"No branch matching '{branch_name}' found in this network.")

    if summary.get("flagship_risk"):
        fr = summary["flagship_risk"]
        match = [b for b in branches if b.get("uninumbr") == fr.get("uninumbr")]
        if match:
            return match[0]
    if summary.get("top5"):
        return summary["top5"][0]
    return branches[0] if branches else None


def get_single_branch_narrative(bank_name, branch, strat, play, vuln_competitors=None):
    """Lightweight narrative call scoped to ONE branch -- deliberately NOT
    the full get_narratives() network-wide schema (exec_headline, priority_focus,
    next_12_months, etc. would be wasted tokens/latency for a single-branch
    preview meant to be fast enough to run live in a pitch meeting). Returns
    the same {"branch_verdicts": ..., "branch_plays": ..., "branch_audiences": ...}
    shape render_branch_deep_dive() already expects, so no changes needed there.

    vuln_competitors: the top-3 vulnerability-ranked competitors (ROA,
    noncurrent %, YoY, vuln_score) for this branch -- i.e. vuln_targets[uninumbr]
    from fetch_vulnerability_targets(). Previously NOT passed here at all: this
    call only ever saw strat['top_competitor'] (name/distance/deposits from the
    adaptive-radius lookup), so the verdict/play/audience text had no visibility
    into which competitors were actually weak, or how this branch's own deposit
    size compares to any of them. That gap -- not the prompt wording -- was the
    main reason the narrative read as thin: it was reasoning from a fraction of
    the data the rendered doc itself displays in the competitor table."""
    branch_label = f"{branch.get('namebr')} ({branch.get('citybr')}, {branch.get('stalpbr')})"
    empty = {"branch_verdicts": {}, "branch_plays": {}, "branch_audiences": {}}

    if not bad.ANTH_KEY or not bad.anthropic:
        return empty

    top_comp = strat.get("top_competitor") if strat else None
    comp_str = (f"nearest named competitor {top_comp.get('bank_name')} "
                f"{bad._sf(top_comp.get('distance_miles')):.2f}mi away with "
                f"${bad._sf(top_comp.get('deposits'))/1e6:.0f}M deposits"
                if top_comp else "no named competitor within the adaptive radius")
    driver_clause = bad._score_driver_clause(branch)
    own_dep = bad._sf(branch.get("latest_dep"))

    vuln_lines = []
    for c in (vuln_competitors or [])[:3]:
        comp_dep = bad._sf(c.get("deposits"))
        ratio_str = ""
        if comp_dep > 0 and own_dep > 0:
            ratio = own_dep / comp_dep
            ratio_str = (f", branch is {ratio:.1f}x their size" if ratio >= 1
                         else f", branch is {ratio:.2f}x their size (smaller)")
        vuln_lines.append(
            f"  #{c.get('rank')} {c.get('bank_name')}: ${comp_dep/1e6:.0f}M deposits, "
            f"{bad._sf(c.get('yoy_pct')):+.1f}% YoY, ROA {bad._sf(c.get('roa')):.2f}%, "
            f"noncurrent {bad._sf(c.get('noncurrent_pct')):.1f}%{ratio_str} — "
            f"{bad._vulnerability_reasoning(c)}"
        )
    vuln_block = ("Vulnerability-ranked named competitors (this branch's real edge over each):\n"
                  + "\n".join(vuln_lines)) if vuln_lines else \
                 "No vulnerability-ranked competitors available for this branch."

    ctx = (
        f"Bank: {bank_name}\n"
        f"Branch: {branch_label}\n"
        f"Score {bad._sf(branch.get('opportunity_score')):.0f}/100, "
        f"zone {branch.get('opportunity_zone')}, "
        f"${own_dep/1e6:.0f}M deposits, "
        f"{bad._sf(branch.get('yoy_deposits'))*100:+.1f}% YoY, {comp_str}."
        + (f" {driver_clause[0].upper()}{driver_clause[1:]}." if driver_clause else "") + "\n"
        f"{vuln_block}\n"
        f"Household income ${bad._sf(branch.get('household_income')):.0f} "
        f"({bad._sf(branch.get('yoy_income_growth'))*100:+.1f}% YoY), "
        f"population YoY {bad._sf(branch.get('yoy_pop_growth'))*100:+.1f}%, "
        f"home value YoY {bad._sf(branch.get('zhvi_yoy_pct')):+.1f}%.\n"
        f"Assigned play: {play or 'n/a'}."
    )

    system = """You are writing a ONE-BRANCH preview of the Verlocity BMAP Assessment,
shown live in a sales pitch to demonstrate analytical depth. Same standards as the full
paid Assessment: confident, commercial, decisive, grounded in the specific numbers given.
Every claim needs a number. Return ONLY valid JSON, no markdown fences:
{
  "branch_verdicts": {"Branch Name (City, ST)": "3-4 sentences synthesizing score, zone, the named competitive threat (or its absence), and deposit trajectory into a clear verdict -- the 'why' behind the assigned play. If a score-driver sentence is given (what's actually driving the score up or down), use it explicitly -- naming the real driver (e.g. 'capped by a shrinking local market, not competition' or 'reflects deposit scale, not underlying growth') is exactly the insight a prospect is paying to see, not a restatement of the number. If vulnerability-ranked competitors are given, the verdict must name at least one specific weakness (declining deposits, weak ROA, elevated noncurrent assets) rather than treating competitors as an undifferentiated group."},
  "branch_plays": {"Branch Name (City, ST)": {"resource_posture": "One sentence, grounded in this branch's specific numbers -- not a generic play-name restatement. If no named competitor exists within the adaptive radius, do NOT write language implying one does. If a relative-size figure is given, use it -- 'this branch is 2.8x the size of its weakest named competitor' is a stronger resourcing argument than restating deposit totals separately.", "media_brief": "One to two sentences, naming the actual target audience and product implied by this branch's specific data. If a specific competitor weakness is given (e.g. a named bank losing deposits or showing balance-sheet stress), the media brief should name conquesting that competitor's depositors as part of the angle, not just describe the branch's own demographics in isolation."}},
  "branch_audiences": {"Branch Name (City, ST)": {
    "narrative": "2-3 sentences using the household income, income YoY, population YoY, home value YoY, AND the competitive weakness data given -- not demographics alone. Frame through Verlocity's AudienceFinder segments (High-Quality Local Prospect, Regression-Scored Lookalike, Competitive Conquesting, Warm Retargeting) where the signal supports it. If a named competitor is losing deposits, Competitive Conquesting targeting THEIR depositor base specifically is a stronger, more concrete angle than generic new-household prospecting.",
    "persona_name": "A short archetype-style label for this branch's dominant audience segment, in the style of 'The Equity Plateau' or 'The Cash-Flow Balancer' -- an income/wealth-stage ARCHETYPE grounded in the actual numbers given. Never a specific named individual (no 'Sarah, 34') -- Verlocity's persona layer names segments, not people.",
    "persona_tagline": "One short line under the persona name (5-10 words) capturing the segment's core motivation.",
    "life_stage": "e.g. 'Mid-career to pre-retirement, 40s-60s' -- inferred from income level and market-maturity signals given, not invented biographical detail.",
    "wealth_signal": "Income range + home value context, grounded directly in the numbers given for this branch.",
    "primary_need": "The banking product or need this segment most likely prioritizes, grounded in the income/growth profile given.",
    "switch_driver": "What would actually move this segment to switch banks or deepen a relationship here -- grounded in the competitive weakness data given, when present.",
    "strong_signals": ["1-2 short (under 15 words) bullets naming concrete reasons this branch's audience opportunity is real -- population growth + timing, income tier + rate sensitivity, flat home appreciation shifting focus to deposit yield. Every bullet needs a number from the data given."],
    "validate_before_activating": ["1 short (under 20 words) honest caveat on what this read can't confirm from the data alone -- e.g. no visibility into existing customer deposit balances or share-of-wallet, or a demographic trend that needs a real-world reason before acting on it. This is what makes the read credible rather than just optimistic -- do not skip it."]
  }}
}

Every weak competitor named in the vulnerability data must be usable material -- do not just describe the #1-ranked competitor and ignore the rest; the verdict, play, or audience sections should collectively reflect that more than one competitor is winnable, when more than one shows real weakness."""

    try:
        client = bad.anthropic.Anthropic(api_key=bad.ANTH_KEY)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1500,
            thinking={"type": "disabled"},
            system=system,
            messages=[{"role": "user", "content": ctx}],
        )
        raw = msg.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        import json
        return json.loads(raw)
    except Exception as e:
        print(f"  ⚠ Single-branch narrative failed ({type(e).__name__}: {e}) — using deterministic fallback")
        return empty


def build_branch_preview_doc(bank_name, branch, strat, play, entry, capped_yoy, narr,
                              total_branch_count, geo_by_uid, tmpdir=".", vuln_targets=None):
    """Short standalone document: intro framing -> the one branch's full deep
    dive (via bad.render_branch_deep_dive, identical to the paid Assessment)
    -> closing CTA. Uses the same brand styling as bmap_assessment_doc.py."""
    cover_title = "BMAP Assessment — Branch Deep Dive Preview"
    cover_subtitle = (f"This is one branch, shown at the exact depth and analytical rigor of the full "
                       f"BMAP Assessment — the same competitive geocoding, adaptive-radius modeling, "
                       f"and capture-dollar sizing your team would receive for every priority branch "
                       f"in {bank_name}'s {total_branch_count}-branch network.")

    # ── Cover + branded header/footer -- prefers Brandon's actual live
    # template (verlocity_cover_template.docx), the SAME asset + mechanism
    # bad.build_assessment_doc() already uses for the $10K Assessment via
    # load_cover_template(). Previously this called bad.build_branded_cover()
    # directly and never tried the template at all, despite the old comment
    # here claiming parity with the Assessment flow -- so every Preview got
    # the placeholder cover (a static flattened image with "Trustmark
    # National Bank" and a fixed branch count baked into the pixels) even
    # in a deploy where the real template was present and the Assessment
    # doc sitting right next to it was already using it correctly. Brandon's
    # template's own default content IS this Preview's own example
    # (Jones Valley/Trustmark) -- load_cover_template() overwrites those
    # runs' text with the real bank name/title/subtitle/date per run and
    # brings the template's own real header/footer with it, so
    # setup_branded_header_footer() is only needed on the fallback path.
    doc = bad.load_cover_template(bank_name, doc_title=cover_title, subtitle=cover_subtitle)
    if doc is not None:
        section = doc.sections[0]
    else:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        bad.build_branded_cover(doc, bank_name, cover_title, subtitle=cover_subtitle)
        bad.setup_branded_header_footer(doc, bank_name)

    # ── The one branch, full depth — identical code path to the paid Assessment ──
    bad.render_branch_deep_dive(
        doc, branch, strat, play, entry, capped_yoy,
        narr.get("branch_verdicts") or {}, narr.get("branch_plays") or {},
        narr.get("branch_audiences") or {}, geo_by_uid,
        tmpdir, heading_space_before=0, vuln_targets=vuln_targets,
    )

    # ── Closing CTA — big number up front, checklist instead of one dense
    # paragraph. The old version was a single wall of bold white text in a
    # box; nothing for the eye to land on first. This leads with the branch
    # count as the actual visual anchor, then lists what's included as
    # scannable lines rather than one run-on sentence. ──
    doc.add_page_break()
    p_cta_h = doc.add_paragraph()
    p_cta_h.paragraph_format.space_before = Pt(20)
    r_cta_h = p_cta_h.add_run("This Is One Branch.")
    r_cta_h.bold = True
    r_cta_h.font.size = Pt(20)
    r_cta_h.font.color.rgb = bad.NAVY
    r_cta_h.font.name = bad.FONT_HEAD

    cta_box = doc.add_table(rows=1, cols=2)
    cta_box.autofit = False
    stat_cell, list_cell = cta_box.rows[0].cells
    stat_cell.width = Inches(1.9)
    list_cell.width = Inches(4.6)
    for c in (stat_cell, list_cell):
        bad._set_cell_shading(c, "083D5F")

    stat_cell.paragraphs[0].text = ""
    stat_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = stat_cell.paragraphs[0].add_run(str(total_branch_count))
    r_num.bold = True
    r_num.font.size = Pt(44)
    r_num.font.color.rgb = RGBColor(0x02, 0xA7, 0xC2)  # teal
    r_num.font.name = bad.FONT_HEAD
    p_num_lbl = stat_cell.add_paragraph()
    p_num_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num_lbl = p_num_lbl.add_run(f"BRANCHES IN\n{bank_name.upper()}'S NETWORK")
    r_num_lbl.font.size = Pt(8)
    r_num_lbl.bold = True
    r_num_lbl.font.color.rgb = RGBColor(0xAF, 0xD8, 0xE2)
    r_num_lbl.font.name = bad.FONT_HEAD

    list_cell.paragraphs[0].text = ""
    p_lead = list_cell.paragraphs[0]
    p_lead.paragraph_format.space_after = Pt(6)
    r_lead = p_lead.add_run("The full BMAP Assessment delivers this exact depth for every priority branch:")
    r_lead.bold = True
    r_lead.font.size = Pt(10.5)
    r_lead.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_lead.font.name = bad.FONT_HEAD

    deliverables = [
        "Verdict + assigned play, per branch",
        "Competitive radius map — all competitors, sized by deposits",
        "Capture-dollar modeling (low / medium / aggressive)",
        "Audience signal, per branch",
        "Network-wide executive synthesis",
        "Live market intelligence + financial benchmarking",
    ]
    for item in deliverables:
        p_item = list_cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(2)
        r_check = p_item.add_run("✓  ")
        r_check.bold = True
        r_check.font.size = Pt(9.5)
        r_check.font.color.rgb = RGBColor(0x02, 0xA7, 0xC2)
        r_check.font.name = bad.FONT_HEAD
        r_item = p_item.add_run(item)
        r_item.font.size = Pt(9.5)
        r_item.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_item.font.name = bad.FONT_HEAD
    list_cell.paragraphs[-1].paragraph_format.space_after = Pt(8)

    # Build stamp -- tiny, gray, easy to ignore, exists so anyone looking at
    # the doc can confirm which code generated it without guessing.
    p_build = doc.add_paragraph()
    p_build.paragraph_format.space_before = Pt(8)
    r_build = p_build.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} · build {bad.GENERATOR_BUILD}")
    r_build.font.size = Pt(7)
    r_build.font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
    r_build.font.name = bad.FONT_HEAD

    return doc



def generate_preview(ik, name_hint=None, branch_name=None, tmpdir="."):
    """Core logic, returns (doc, bank_name, branch_label, total_branch_count)
    or raises ValueError with a clear message. Separated from run() so both
    the CLI entry point (saves to disk) and the Hub's Quick Export endpoint
    (streams from memory, same pattern as /generate-assessment) can share
    it without duplicating the fetch/pick/narrate/build sequence.

    Uses skip_competitive_strategy=True + fetch_single_branch_strategy() for
    the target branch specifically, rather than the full $10K Assessment's
    network-wide competitive-strategy fetch. That network-wide fetch pages
    through every branch's competitor rows (11,000+ for a 166-branch network
    like Trustmark) and this flow only ever used one branch's slice of it --
    real cause of a production timeout once that batch call started being
    correctly paginated instead of silently truncated. This keeps the
    Preview fast regardless of network size, which matters since it's meant
    to run live in a pitch meeting."""
    print(f"[branch-preview] fetching network data for {ik}...")
    d = bad.fetch_full_network_data(ik, skip_competitive_strategy=True)
    if not d["branches"]:
        raise ValueError(f"No branch data found for inst_key='{ik}'. This bank may not be "
                          f"ingested into BMAP yet, or the inst_key is incorrect.")

    bank_name = name_hint or (d["branches"][0].get("namefull") if d["branches"] else None) or ik
    branches = d["branches"]
    capped_yoy = d.get("capped_yoy") or {}
    vulnerability_targets = d.get("vulnerability_targets") or {}
    geo_by_uid = {g["uninumbr"]: g for g in (d.get("branches_geo") or []) if g.get("uninumbr") is not None}

    summary = bad.summarize_network(d)

    target = pick_preview_branch(summary, branches, branch_name)
    if not target:
        raise ValueError("Could not select a branch to preview — network has no branches.")

    target_geo = geo_by_uid.get(target.get("uninumbr")) or {}
    target_with_geo = {**target, "lat": target_geo.get("lat"), "lon": target_geo.get("lon")}
    strat = bad.fetch_single_branch_strategy(ik, target_with_geo)
    play = bad.get_play(target.get("opportunity_zone"), target.get("matrix_quadrant"))

    # Capture pool must come from the SAME competitor named as the priority
    # target in the narrative (vulnerability-ranked, what the assigned play
    # actually targets) -- not strat["top_competitor"] (adaptive-radius,
    # largest-nearby-deposit competitor, a different selection entirely).
    # This was the Marc Winkler bug: capture math ran against Wells Fargo's
    # deposits while the recommended play targeted United Community Bank /
    # Bryant Bank. Falls back to top_competitor only when there's no
    # vulnerability-ranked target, matching build_branch_deep_dives().
    branch_vuln_list = sorted(
        (vulnerability_targets or {}).get(target.get("uninumbr"), []),
        key=lambda c: c.get("rank") or 99
    )
    top_comp = strat.get("top_competitor") if strat else None
    capture_target = branch_vuln_list[0] if branch_vuln_list else top_comp
    capture_pool = bad._sf(capture_target.get("deposits")) if capture_target else 0.0
    entry = {"branch": target, "strategy": strat, "play": play, "capture_pool": capture_pool}

    branch, strat, play = entry["branch"], entry["strategy"], entry["play"]
    branch_label = f"{branch.get('namebr')} ({branch.get('citybr')}, {branch.get('stalpbr')})"
    print(f"[branch-preview] previewing {branch_label}")
    narr = get_single_branch_narrative(bank_name, branch, strat, play, branch_vuln_list)
    doc = build_branch_preview_doc(bank_name, branch, strat, play, entry, capped_yoy, narr,
                                    len(branches), geo_by_uid, tmpdir=tmpdir,
                                    vuln_targets=vulnerability_targets)
    return doc, bank_name, branch_label, len(branches)


def run(ik, name_hint=None, branch_name=None):
    import tempfile
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            doc, bank_name, branch_label, total = generate_preview(ik, name_hint, branch_name, tmpdir)
        except ValueError as e:
            print(f"[branch-preview] ✗ {e}")
            return None

        safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in bank_name).strip()
        date = datetime.now().strftime("%Y%m%d")
        branch_safe = "".join(c if c.isalnum() or c in " _-" else "_"
                               for c in branch_label.split(" (")[0]).strip()
        filename = f"BMAP_Preview_{safe}_{branch_safe}_{date}.docx"
        out_path = os.path.join(bad.OUT_DIR if hasattr(bad, "OUT_DIR") else ".", filename)
        doc.save(out_path)

    print(f"[branch-preview] ✓ {out_path}")
    return out_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a single-branch BMAP Assessment preview")
    parser.add_argument("--inst_key", required=True, help="e.g. bank_463735")
    parser.add_argument("--name", default=None, help="Bank display name (optional)")
    parser.add_argument("--branch", default=None,
                         help="Branch name to preview (optional -- auto-selects the most "
                              "compelling branch if omitted: flagship risk, else top opportunity score)")
    args = parser.parse_args()
    result = run(args.inst_key, args.name, args.branch)
    sys.exit(0 if result else 1)
